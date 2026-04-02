import os
import re
import json
import sqlite3
import shutil
from datetime import datetime
from pathlib import Path

from flask import Flask, render_template, request, jsonify, send_from_directory, abort
from docx import Document
from docx.oxml.ns import qn

from functools import wraps
import hashlib

app = Flask(__name__)
app.secret_key = "ulteria-gestionale-2026-secret-key"

BASE_DIR = Path(__file__).resolve().parent
UPLOADS_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "output"
CONFIG_PATH = BASE_DIR / "config.json"
DB_PATH = BASE_DIR / "database.db"

TEMPLATE_MAP = {
    "E40": "«NR»_ACCORDO QUADRO_«STU»_SOSTITUZIONE RIPARTITORI CALORE OMS_E40.docx",
    "Q55": "«NR»_ACCORDO QUADRO_«STU»_SOSTITUZIONE RIPARTITORI CALORE OMS_Q5.55.docx",
}


# ── Database ─────────────────────────────────────────────────────────

def get_db():
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA foreign_keys=ON")
    return conn


def init_db():
    conn = get_db()
    conn.executescript("""
        /* ── Step 1: Etichette dinamiche ── */
        CREATE TABLE IF NOT EXISTS etichette (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            categoria TEXT NOT NULL,
            valore TEXT NOT NULL,
            colore_bg TEXT DEFAULT '#F4F9FD',
            colore_testo TEXT DEFAULT '#0D1F35',
            ordine INTEGER DEFAULT 0,
            attiva INTEGER DEFAULT 1,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        );

        /* ── Agenti ── */
        CREATE TABLE IF NOT EXISTS agenti (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT, cognome TEXT, email TEXT, telefono TEXT,
            colore TEXT DEFAULT '#009FE3', note TEXT,
            data_inserimento DATETIME
        );

        /* ── Clienti ── */
        CREATE TABLE IF NOT EXISTS clienti (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome_studio TEXT UNIQUE, via TEXT, cap TEXT, citta TEXT,
            email TEXT, telefono TEXT, referente TEXT, note TEXT,
            tipo_cliente TEXT DEFAULT 'Amministratore',
            settore TEXT, note_generali TEXT,
            data_inserimento DATETIME
        );

        /* ── Step 2: Oggetti (condomini/edifici/servizi) ── */
        CREATE TABLE IF NOT EXISTS oggetti (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            cliente_id INTEGER NOT NULL,
            nome TEXT, via TEXT NOT NULL, civico TEXT,
            comune TEXT NOT NULL, provincia TEXT, cap TEXT,
            tipo_oggetto TEXT DEFAULT 'condominio',
            stato_pipeline TEXT DEFAULT 'prospect',
            agente_id INTEGER, natura TEXT,
            n_unita INTEGER, n_scale INTEGER,
            cliente_precedente_id INTEGER,
            data_cambio_intestazione DATE,
            data_primo_contatto DATE, data_ultimo_contatto DATE,
            motivo_perdita TEXT, note_perdita TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (cliente_id) REFERENCES clienti(id),
            FOREIGN KEY (agente_id) REFERENCES agenti(id)
        );

        /* ── Offerte ── */
        CREATE TABLE IF NOT EXISTS offerte (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            numero INTEGER, nome_studio TEXT, nome_condominio TEXT,
            via TEXT, cap TEXT, citta TEXT, riferimento TEXT, template TEXT,
            prezzo_fornitura REAL, prezzo_care REAL, canone_lettura REAL,
            modalita TEXT DEFAULT 'vendita', totale REAL,
            stato TEXT DEFAULT 'richiamato', email_studio TEXT,
            data_creazione DATETIME, path_docx TEXT, path_pdf TEXT, note TEXT,
            agente_id INTEGER REFERENCES agenti(id),
            motivo_perdita TEXT, note_perdita TEXT, importo REAL,
            condominio_id INTEGER, oggetto_id INTEGER,
            versione TEXT DEFAULT 'A',
            offerta_padre_id INTEGER REFERENCES offerte(id),
            stato_versione TEXT DEFAULT 'attiva',
            tipo_offerta TEXT DEFAULT 'installazione',
            natura TEXT DEFAULT 'nuovo',
            importo_servizio_annuo REAL,
            is_accordo_quadro INTEGER DEFAULT 0
        );

        /* ── Step 3: Righe economiche offerta ── */
        CREATE TABLE IF NOT EXISTS offerte_righe (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            offerta_id INTEGER NOT NULL,
            descrizione TEXT NOT NULL,
            tipo_riga TEXT DEFAULT 'fornitura',
            prezzo_unitario REAL, quantita REAL,
            quantita_stimata INTEGER DEFAULT 0,
            totale_riga REAL, ordine INTEGER DEFAULT 0,
            FOREIGN KEY (offerta_id) REFERENCES offerte(id)
        );

        /* ── Step 4: Timeline eventi ── */
        CREATE TABLE IF NOT EXISTS timeline_eventi (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            tipo_evento TEXT NOT NULL,
            descrizione TEXT NOT NULL,
            cliente_id INTEGER, oggetto_id INTEGER, offerta_id INTEGER,
            utente TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        );

        /* ── Step 5: Attivita ── */
        CREATE TABLE IF NOT EXISTS attivita (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            agente_id INTEGER NOT NULL,
            tipo TEXT NOT NULL DEFAULT 'todo',
            titolo TEXT NOT NULL, descrizione TEXT,
            data_scadenza DATETIME, priorita TEXT DEFAULT 'media',
            stato TEXT DEFAULT 'aperta',
            cliente_id INTEGER, oggetto_id INTEGER, offerta_id INTEGER,
            completato_il DATETIME,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (agente_id) REFERENCES agenti(id),
            FOREIGN KEY (cliente_id) REFERENCES clienti(id),
            FOREIGN KEY (offerta_id) REFERENCES offerte(id)
        );

        /* ── Step 6: Users ── */
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            nome TEXT NOT NULL, cognome TEXT NOT NULL,
            ruolo TEXT NOT NULL DEFAULT 'agente',
            agente_id INTEGER REFERENCES agenti(id),
            is_active INTEGER DEFAULT 1,
            last_login DATETIME,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        );

        /* ── Note (extended) ── */
        CREATE TABLE IF NOT EXISTS note_clienti (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            cliente_id INTEGER, oggetto_id INTEGER, offerta_id INTEGER,
            testo TEXT NOT NULL, autore TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (cliente_id) REFERENCES clienti(id)
        );

        /* ── Legacy compat: condomini ── */
        CREATE TABLE IF NOT EXISTS condomini (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            cliente_id INTEGER NOT NULL, nome TEXT NOT NULL,
            indirizzo TEXT, comune TEXT, provincia TEXT,
            stato_pipeline TEXT DEFAULT 'prospect',
            agente_id INTEGER, data_assemblea DATE,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
        );

        /* ── Prompt 1: Prodotti (listino prezzi) ── */
        CREATE TABLE IF NOT EXISTS prodotti (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            codice TEXT UNIQUE NOT NULL,
            nome TEXT NOT NULL,
            categoria TEXT NOT NULL,
            modello TEXT,
            trasmissione TEXT,
            dn TEXT,
            prezzo_acquisto REAL,
            prezzo_vendita REAL,
            data_ultimo_prezzo DATE,
            fornitore TEXT,
            note TEXT,
            attivo INTEGER DEFAULT 1,
            ordine INTEGER DEFAULT 0,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
        );

        /* ── Prompt 1: Modelli apparecchio ── */
        CREATE TABLE IF NOT EXISTS modelli_apparecchio (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            categoria TEXT NOT NULL,
            nome TEXT NOT NULL,
            icona TEXT DEFAULT 'thermometer',
            trasmissione_disponibile INTEGER DEFAULT 0,
            dn_disponibile INTEGER DEFAULT 0,
            tipo_lettura TEXT,
            attivo INTEGER DEFAULT 1,
            ordine INTEGER DEFAULT 0,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        );

        /* ── Prompt 1: Prezzi installazione base ── */
        CREATE TABLE IF NOT EXISTS prezzi_installazione (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            tipo TEXT NOT NULL UNIQUE,
            descrizione TEXT NOT NULL,
            prezzo_base REAL NOT NULL DEFAULT 0,
            unita TEXT DEFAULT 'cad',
            note TEXT,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
        );

        /* ── Prompt 1: Foglio costi per oggetto ── */
        CREATE TABLE IF NOT EXISTS fogli_costi (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            oggetto_id INTEGER NOT NULL,
            offerta_id INTEGER,
            costo_apparecchi REAL DEFAULT 0,
            costo_installazione_idraulica REAL DEFAULT 0,
            costo_installazione_elettrica REAL DEFAULT 0,
            costo_concentratori REAL DEFAULT 0,
            costo_materiali_extra REAL DEFAULT 0,
            note_costi TEXT,
            totale_costi REAL DEFAULT 0,
            ricavo_fornitura REAL DEFAULT 0,
            ricavo_servizio_annuo REAL DEFAULT 0,
            k_moltiplicatore REAL DEFAULT 1.0,
            margine_euro REAL DEFAULT 0,
            margine_percentuale REAL DEFAULT 0,
            provvigione1_nome TEXT,
            provvigione1_percentuale REAL DEFAULT 0,
            provvigione1_euro REAL DEFAULT 0,
            provvigione2_nome TEXT,
            provvigione2_percentuale REAL DEFAULT 0,
            provvigione2_euro REAL DEFAULT 0,
            provvigione3_nome TEXT,
            provvigione3_percentuale REAL DEFAULT 0,
            provvigione3_euro REAL DEFAULT 0,
            netto_finale REAL DEFAULT 0,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (oggetto_id) REFERENCES oggetti(id),
            FOREIGN KEY (offerta_id) REFERENCES offerte(id)
        );

        /* ── Prompt 1: Righe extra foglio costi ── */
        CREATE TABLE IF NOT EXISTS foglio_costi_extra (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            foglio_costi_id INTEGER NOT NULL,
            descrizione TEXT NOT NULL,
            quantita REAL DEFAULT 1,
            prezzo_unitario REAL DEFAULT 0,
            totale REAL DEFAULT 0,
            FOREIGN KEY (foglio_costi_id) REFERENCES fogli_costi(id)
        );

        /* ── Prompt 1: Storico prezzi prodotti ── */
        CREATE TABLE IF NOT EXISTS prodotti_prezzi_storico (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            prodotto_id INTEGER NOT NULL,
            prezzo_acquisto_precedente REAL,
            prezzo_vendita_precedente REAL,
            prezzo_acquisto_nuovo REAL,
            prezzo_vendita_nuovo REAL,
            aggiornato_da TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (prodotto_id) REFERENCES prodotti(id)
        );

        /* ── Prompt 4: Segnalatori ── */
        CREATE TABLE IF NOT EXISTS segnalatori (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL,
            tipo TEXT NOT NULL DEFAULT 'segnalatore',
            azienda TEXT, email TEXT, telefono TEXT,
            provvigione_default_pct REAL DEFAULT 0,
            note TEXT, attivo INTEGER DEFAULT 1,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        );
        CREATE TABLE IF NOT EXISTS offerte_segnalatori (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            offerta_id INTEGER NOT NULL,
            segnalatore_id INTEGER NOT NULL,
            provvigione_pct REAL DEFAULT 0,
            importo_base REAL DEFAULT 0,
            provvigione_euro REAL DEFAULT 0,
            stato_pagamento TEXT DEFAULT 'da_pagare',
            data_pagamento DATE, note TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (offerta_id) REFERENCES offerte(id),
            FOREIGN KEY (segnalatore_id) REFERENCES segnalatori(id)
        );
        CREATE TABLE IF NOT EXISTS oggetti_agenti_storico (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            oggetto_id INTEGER NOT NULL,
            agente_id_precedente INTEGER,
            agente_id_nuovo INTEGER NOT NULL,
            motivo TEXT, data_cambio DATE NOT NULL,
            effettuato_da TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (oggetto_id) REFERENCES oggetti(id)
        );
    """)

    # ── Step 7: Migrations for existing DBs ──
    migrations = [
        "ALTER TABLE offerte ADD COLUMN agente_id INTEGER REFERENCES agenti(id)",
        "ALTER TABLE agenti ADD COLUMN colore TEXT DEFAULT '#009FE3'",
        "ALTER TABLE offerte ADD COLUMN motivo_perdita TEXT",
        "ALTER TABLE offerte ADD COLUMN note_perdita TEXT",
        "ALTER TABLE offerte ADD COLUMN importo REAL",
        "ALTER TABLE offerte ADD COLUMN condominio_id INTEGER",
        "ALTER TABLE offerte ADD COLUMN oggetto_id INTEGER",
        "ALTER TABLE offerte ADD COLUMN versione TEXT DEFAULT 'A'",
        "ALTER TABLE offerte ADD COLUMN offerta_padre_id INTEGER",
        "ALTER TABLE offerte ADD COLUMN stato_versione TEXT DEFAULT 'attiva'",
        "ALTER TABLE offerte ADD COLUMN tipo_offerta TEXT DEFAULT 'installazione'",
        "ALTER TABLE offerte ADD COLUMN natura TEXT DEFAULT 'nuovo'",
        "ALTER TABLE offerte ADD COLUMN importo_servizio_annuo REAL",
        "ALTER TABLE offerte ADD COLUMN is_accordo_quadro INTEGER DEFAULT 0",
        "ALTER TABLE clienti ADD COLUMN tipo_cliente TEXT DEFAULT 'Amministratore'",
        "ALTER TABLE clienti ADD COLUMN settore TEXT",
        "ALTER TABLE clienti ADD COLUMN note_generali TEXT",
        "ALTER TABLE attivita ADD COLUMN oggetto_id INTEGER",
        "ALTER TABLE note_clienti ADD COLUMN oggetto_id INTEGER",
        "ALTER TABLE note_clienti ADD COLUMN offerta_id INTEGER",
        # ── Prompt 3B: Foglio Costi unificato ──
        "ALTER TABLE fogli_costi ADD COLUMN scenario TEXT DEFAULT 'valvole'",
        "ALTER TABLE fogli_costi ADD COLUMN installatore_idraulico TEXT",
        "ALTER TABLE fogli_costi ADD COLUMN cont_riscaldamento INTEGER DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN cont_riscaldamento_trasmissione TEXT",
        "ALTER TABLE fogli_costi ADD COLUMN cont_riscaldamento_dn INTEGER",
        "ALTER TABLE fogli_costi ADD COLUMN cont_riscaldamento_costo REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN cont_hc INTEGER DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN cont_hc_trasmissione TEXT",
        "ALTER TABLE fogli_costi ADD COLUMN cont_hc_dn INTEGER",
        "ALTER TABLE fogli_costi ADD COLUMN cont_hc_costo REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN cont_raffrescamento INTEGER DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN cont_raffrescamento_trasmissione TEXT",
        "ALTER TABLE fogli_costi ADD COLUMN cont_raffrescamento_dn INTEGER",
        "ALTER TABLE fogli_costi ADD COLUMN cont_raffrescamento_costo REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN cont_acqua_calda INTEGER DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN cont_acqua_calda_trasmissione TEXT",
        "ALTER TABLE fogli_costi ADD COLUMN cont_acqua_calda_dn INTEGER",
        "ALTER TABLE fogli_costi ADD COLUMN cont_acqua_calda_costo REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN cont_acqua_fredda INTEGER DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN cont_acqua_fredda_trasmissione TEXT",
        "ALTER TABLE fogli_costi ADD COLUMN cont_acqua_fredda_dn INTEGER",
        "ALTER TABLE fogli_costi ADD COLUMN cont_acqua_fredda_costo REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN cont_acqua_ricircolo INTEGER DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN cont_acqua_ricircolo_costo REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN cont_acqua_duale INTEGER DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN cont_acqua_duale_costo REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN inst_cont_calore REAL DEFAULT 27",
        "ALTER TABLE fogli_costi ADD COLUMN inst_cont_acqua_calda REAL DEFAULT 22",
        "ALTER TABLE fogli_costi ADD COLUMN inst_cont_acqua_fredda REAL DEFAULT 22",
        "ALTER TABLE fogli_costi ADD COLUMN inst_modifiche_idrauliche REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN costo_valvola_zona REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN costo_attuatore REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN costo_produzione_modulo REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN costo_opere_idrauliche_extra REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN costo_trasformatore REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN costo_rele REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN costo_elettricista REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN costo_collegamenti_elettrici REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN costo_valvola_intercettazione REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN centr_famiglia TEXT",
        "ALTER TABLE fogli_costi ADD COLUMN centr_modello TEXT",
        "ALTER TABLE fogli_costi ADD COLUMN centr_pezzi INTEGER DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN centr_costo_acquisto REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN centr_costo_installazione REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN centr_pw_router REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN centr_collaudo TEXT DEFAULT 'incluso'",
        "ALTER TABLE fogli_costi ADD COLUMN n_radiatori INTEGER DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN costo_kit_valvola REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN costo_montaggio_valvola REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN costo_apparecchio_ripartitore REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN costo_extra_trasporto REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN servizio_lettura_tipo TEXT",
        "ALTER TABLE fogli_costi ADD COLUMN servizio_lettura_cad REAL DEFAULT 0",
        "ALTER TABLE fogli_costi ADD COLUMN servizio_lettura_totale REAL DEFAULT 0",
        # ── Prompt 4 migrations ──
        "ALTER TABLE offerte ADD COLUMN segnalatore_id INTEGER REFERENCES segnalatori(id)",
        "ALTER TABLE offerte ADD COLUMN tipo_servizio TEXT",
        "ALTER TABLE offerte ADD COLUMN tipo_servizio_secondario TEXT",
        # ── Prompt 5: Riepilogo refactoring ──
        "ALTER TABLE offerte ADD COLUMN macro_categoria TEXT",
        "ALTER TABLE offerte ADD COLUMN sottotipo TEXT",
        "ALTER TABLE offerte ADD COLUMN is_gara_appalto INTEGER DEFAULT 0",
        "ALTER TABLE offerte ADD COLUMN valore_gara REAL",
        "ALTER TABLE offerte ADD COLUMN gara_id TEXT",
        "ALTER TABLE offerte ADD COLUMN valore_commessa REAL",
    ]
    for stmt in migrations:
        try:
            conn.execute(stmt)
        except Exception:
            pass

    # ── Prompt 3B: Add M-BUS concentrators ──
    mbus_models = [
        ("concentratore_mbus", "WEBLOG-250", "server", 0, 0, None, 11),
        ("concentratore_mbus", "CME", "server", 0, 0, None, 12),
        ("concentratore_mbus", "CMC", "server", 0, 0, None, 13),
    ]
    for cat, nome, ico, trasm, dn, tl, ordine in mbus_models:
        try:
            conn.execute(
                "INSERT INTO modelli_apparecchio (categoria,nome,icona,trasmissione_disponibile,dn_disponibile,tipo_lettura,ordine) VALUES (?,?,?,?,?,?,?)",
                (cat, nome, ico, trasm, dn, tl, ordine),
            )
        except Exception:
            pass

    # ── Insert default etichette if empty ──
    count = conn.execute("SELECT COUNT(*) FROM etichette").fetchone()[0]
    if count == 0:
        defaults = [
            ("tipo_cliente", "Amministratore", "#E6F5FC", "#0080B8"),
            ("tipo_cliente", "Gestore", "#EAF3DE", "#639922"),
            ("tipo_cliente", "Costruttore", "#FAEEDA", "#854F0B"),
            ("tipo_cliente", "Progettista", "#EEEDFE", "#534AB7"),
            ("tipo_cliente", "Condomino", "#F1EFE8", "#5F5E5A"),
            ("tipo_cliente", "Rivenditore", "#FCE4EC", "#880E4F"),
            ("stato_pipeline", "prospect", "#E6F5FC", "#0080B8"),
            ("stato_pipeline", "offerta_inviata", "#FAEEDA", "#854F0B"),
            ("stato_pipeline", "in_attesa_assemblea", "#FFF3E0", "#E65100"),
            ("stato_pipeline", "preso", "#EAF3DE", "#639922"),
            ("stato_pipeline", "perso", "#FCEBEB", "#A32D2D"),
            ("stato_pipeline", "rimandato", "#EEEDFE", "#534AB7"),
            ("tipo_attivita", "chiamata", "#E6F5FC", "#0080B8"),
            ("tipo_attivita", "email", "#EAF3DE", "#639922"),
            ("tipo_attivita", "visita", "#FAEEDA", "#854F0B"),
            ("tipo_attivita", "assemblea", "#EEEDFE", "#534AB7"),
            ("tipo_attivita", "to-do", "#F1EFE8", "#5F5E5A"),
            ("tipo_attivita", "altro", "#F4F9FD", "#0D1F35"),
            ("motivo_perdita", "Prezzo", "#FCEBEB", "#A32D2D"),
            ("motivo_perdita", "Competitor", "#FCEBEB", "#A32D2D"),
            ("motivo_perdita", "Assemblea non approva", "#FCEBEB", "#A32D2D"),
            ("motivo_perdita", "Non risponde", "#FCEBEB", "#A32D2D"),
            ("motivo_perdita", "Rimandato", "#EEEDFE", "#534AB7"),
            ("motivo_perdita", "Cambio amministratore", "#FAEEDA", "#854F0B"),
            ("motivo_perdita", "Altro", "#F1EFE8", "#5F5E5A"),
            ("tipo_offerta", "fornitura", "#E6F5FC", "#0080B8"),
            ("tipo_offerta", "installazione", "#EAF3DE", "#639922"),
            ("tipo_offerta", "servizio", "#FAEEDA", "#854F0B"),
            ("settore", "amministratori", "#E6F5FC", "#0080B8"),
            ("settore", "progettisti", "#EEEDFE", "#534AB7"),
            ("settore", "gestori", "#EAF3DE", "#639922"),
            ("settore", "costruttori", "#FAEEDA", "#854F0B"),
            # Prompt 4: tipo_servizio
            ("tipo_servizio", "RK", "#E6F5FC", "#0080B8"),
            ("tipo_servizio", "RD", "#EAF3DE", "#639922"),
            ("tipo_servizio", "MANSIS", "#FAEEDA", "#854F0B"),
            ("tipo_servizio", "MANCT", "#EEEDFE", "#534AB7"),
            # Prompt 5: macro_categoria
            ("macro_categoria", "installazione", "#E6F5FC", "#0080B8"),
            ("macro_categoria", "servizi", "#EAF3DE", "#3B6D11"),
            ("macro_categoria", "cc_modus", "#FAEEDA", "#854F0B"),
            ("macro_categoria", "cu_unitron", "#EEEDFE", "#3C3489"),
            ("macro_categoria", "fornitura", "#FAECE7", "#993C1D"),
            ("macro_categoria", "interventi", "#F1EFE8", "#5F5E5A"),
            # Prompt 5: sottotipo
            ("sottotipo", "CK", "#E6F5FC", "#0080B8"),
            ("sottotipo", "CL", "#B5D4F4", "#042C53"),
            ("sottotipo", "RK", "#EAF3DE", "#3B6D11"),
            ("sottotipo", "RD", "#C0DD97", "#173404"),
            ("sottotipo", "MANSIS", "#FAEEDA", "#854F0B"),
            ("sottotipo", "MAN-DOMO", "#FAC775", "#412402"),
            ("sottotipo", "SIM-HOSTING", "#F1EFE8", "#5F5E5A"),
            ("sottotipo", "MANCT", "#EEEDFE", "#3C3489"),
            ("sottotipo", "AVV-SIS", "#FAECE7", "#993C1D"),
            ("sottotipo", "AVV-CONT-MBUS", "#F5C4B3", "#4A1B0C"),
            ("sottotipo", "AVV-RADIO", "#F0997B", "#4A1B0C"),
            ("sottotipo", "MISURATORI", "#FAECE7", "#993C1D"),
            ("sottotipo", "RICAMBI", "#F5C4B3", "#4A1B0C"),
            ("sottotipo", "CM", "#F1EFE8", "#5F5E5A"),
        ]
        for cat, val, bg, txt in defaults:
            conn.execute(
                "INSERT INTO etichette (categoria,valore,colore_bg,colore_testo) VALUES (?,?,?,?)",
                (cat, val, bg, txt),
            )

    # ── Prompt 4: Default segnalatori ──
    sc = conn.execute("SELECT COUNT(*) FROM segnalatori").fetchone()[0]
    if sc == 0:
        conn.execute("INSERT INTO segnalatori (nome,tipo,provvigione_default_pct) VALUES (?,?,?)", ("Piaggi", "manutentore", 0))
        conn.execute("INSERT INTO segnalatori (nome,tipo,provvigione_default_pct) VALUES (?,?,?)", ("Merlotti", "segnalatore", 0))

    # ── Prompt 1: Insert default modelli apparecchio ──
    mc = conn.execute("SELECT COUNT(*) FROM modelli_apparecchio").fetchone()[0]
    if mc == 0:
        modelli_def = [
            ("ripartitore", "E-ITN40", "thermometer", 0, 0, "RK", 1),
            ("ripartitore", "Q5.5", "thermometer", 0, 0, "RK", 2),
            ("contatore_acqua", "SMART-WB", "droplets", 1, 1, "RK", 3),
            ("contatore_acqua", "CSU-R", "droplets", 1, 1, "RK", 4),
            ("contatore_calore", "ELF2", "flame", 1, 0, "RD", 5),
            ("contatore_calore", "H5-HC", "flame", 1, 0, "RD", 6),
            ("concentratore", "HUB-GATEWAY", "wifi", 0, 0, None, 7),
            ("concentratore", "Q-DIRECT", "wifi", 0, 0, None, 8),
            ("concentratore", "Q-NODE", "wifi", 0, 0, None, 9),
            ("concentratore", "Q-GATEWAY", "wifi", 0, 0, None, 10),
        ]
        for cat, nome, ico, trasm, dn, tl, ordine in modelli_def:
            conn.execute(
                "INSERT INTO modelli_apparecchio (categoria,nome,icona,trasmissione_disponibile,dn_disponibile,tipo_lettura,ordine) VALUES (?,?,?,?,?,?,?)",
                (cat, nome, ico, trasm, dn, tl, ordine),
            )

    # ── Prompt 1: Insert default prezzi installazione ──
    pc = conn.execute("SELECT COUNT(*) FROM prezzi_installazione").fetchone()[0]
    if pc == 0:
        prezzi_def = [
            ("ripartitore", "Installazione ripartitore di calore", 0, "cad"),
            ("contatore_acqua", "Installazione contatore acqua", 20, "cad"),
            ("contatore_calore", "Installazione contatore calore", 30, "cad"),
            ("concentratore", "Installazione concentratore/gateway", 0, "cad"),
            ("elettrica_base", "Installazione elettrica base", 0, "corpo"),
        ]
        for tipo, desc, prezzo, unita in prezzi_def:
            conn.execute(
                "INSERT INTO prezzi_installazione (tipo,descrizione,prezzo_base,unita) VALUES (?,?,?,?)",
                (tipo, desc, prezzo, unita),
            )

    conn.commit()
    conn.close()


# ── Config ───────────────────────────────────────────────────────────

def read_config():
    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def write_config(data):
    with open(CONFIG_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


# ── Word helpers ─────────────────────────────────────────────────────

def replace_runs(paragraphs, old, new):
    for p in paragraphs:
        for r in p.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)


def replace_xml(element, old, new):
    for t in element.iter(qn("w:t")):
        if t.text and old in t.text:
            t.text = t.text.replace(old, new)


def fix_stu_via_paragraph(doc, nome_studio, via_studio):
    """Rewrite STU/VIA paragraph: remove all tab-based layout, create two
    clean right-aligned lines using w:br (line break) in Word XML."""
    from lxml import etree
    from copy import deepcopy

    for p in doc.paragraphs:
        has_stu = any(nome_studio in r.text for r in p.runs)
        has_via = any(via_studio in r.text for r in p.runs)
        if not (has_stu and has_via):
            continue

        # Change paragraph alignment from JUSTIFY to RIGHT
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None:
            jc = pPr.find(qn("w:jc"))
            if jc is not None:
                jc.set(qn("w:val"), "right")
            else:
                jc = etree.SubElement(pPr, qn("w:jc"))
                jc.set(qn("w:val"), "right")
            # Remove first-line indent which shifts text
            ind = pPr.find(qn("w:ind"))
            if ind is not None:
                for attr in ["firstLine", "left", "hanging"]:
                    k = qn("w:" + attr)
                    if k in ind.attrib:
                        del ind.attrib[k]

        # Collect run info: find STU run and VIA run
        stu_run_el = None
        via_run_el = None
        tab_runs = []
        for r in p._element.findall(qn("w:r")):
            texts = r.findall(qn("w:t"))
            tabs = r.findall(qn("w:tab"))
            text_content = "".join(t.text or "" for t in texts)
            if nome_studio in text_content:
                stu_run_el = r
            elif via_studio in text_content:
                via_run_el = r
            elif tabs and not text_content.strip():
                tab_runs.append(r)

        if stu_run_el is None or via_run_el is None:
            break

        # Remove ALL tab-only runs from paragraph
        for tr in tab_runs:
            p._element.remove(tr)

        # Insert a line break run between STU and VIA
        # Create a run with <w:br/> + tabs for indentation
        br_run = deepcopy(stu_run_el)
        # Clear text and tab elements from the copy
        for child in list(br_run):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag in ("t", "tab"):
                br_run.remove(child)
        # Add line break element
        rPr = br_run.find(qn("w:rPr"))
        insert_pos = list(br_run).index(rPr) + 1 if rPr is not None else 0
        br_el = etree.Element(qn("w:br"))
        br_run.insert(insert_pos, br_el)

        # Insert br_run right before via_run_el
        p._element.insert(list(p._element).index(via_run_el), br_run)

        # Copy formatting from STU run to VIA run (but keep non-bold)
        via_rPr = via_run_el.find(qn("w:rPr"))
        if via_rPr is None:
            via_rPr = etree.SubElement(via_run_el, qn("w:rPr"))
            via_run_el.insert(0, via_rPr)

        break


def studio_slug(name):
    words = re.sub(r"[^A-Za-z0-9 ]", "", name).split()
    slug = "_".join(words[:3]).upper()
    return slug[:30]


def format_eur(val):
    if val is None:
        return "—"
    return f"{val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


# ── Step 10: Auth ────────────────────────────────────────────────────

def hash_pw(pw):
    return hashlib.sha256(pw.encode()).hexdigest()


def get_current_user():
    from flask import session
    uid = session.get("user_id")
    if not uid:
        return None
    conn = get_db()
    user = conn.execute("SELECT * FROM users WHERE id=? AND is_active=1", (uid,)).fetchone()
    conn.close()
    return dict(user) if user else None


def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        user = get_current_user()
        if not user:
            from flask import redirect, url_for
            return redirect(url_for("login_page"))
        return f(*args, **kwargs)
    return decorated


def require_role(*roles):
    def decorator(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            user = get_current_user()
            if not user:
                return redirect(url_for("login_page"))
            if user["ruolo"] not in roles:
                abort(403)
            return f(*args, **kwargs)
        return decorated
    return decorator


@app.route("/login", methods=["GET", "POST"])
def login_page():
    from flask import session, redirect
    if request.method == "GET":
        return render_template("login.html")
    data = request.form
    email = data.get("email", "")
    pw = data.get("password", "")
    conn = get_db()
    user = conn.execute(
        "SELECT * FROM users WHERE email=? AND is_active=1", (email,)
    ).fetchone()
    conn.close()
    if user and user["password_hash"] == hash_pw(pw):
        session["user_id"] = user["id"]
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        conn = get_db()
        conn.execute("UPDATE users SET last_login=? WHERE id=?", (now, user["id"]))
        conn.commit()
        conn.close()
        return redirect("/")
    return render_template("login.html", error="Email o password non validi")


@app.route("/logout")
def logout_page():
    from flask import session, redirect
    session.clear()
    return redirect("/login")


@app.route("/api/auth/me", methods=["GET"])
def api_auth_me():
    user = get_current_user()
    if not user:
        return jsonify({"ok": False, "error": "Non autenticato"}), 401
    return jsonify({"ok": True, "data": {
        "id": user["id"], "email": user["email"], "nome": user["nome"],
        "cognome": user["cognome"], "ruolo": user["ruolo"], "agente_id": user["agente_id"],
    }})


# ── Step 10: Users API ──────────────────────────────────────────────

@app.route("/api/users", methods=["GET"])
def api_users_list():
    conn = get_db()
    rows = conn.execute("SELECT id,email,nome,cognome,ruolo,agente_id,is_active,last_login FROM users ORDER BY nome").fetchall()
    conn.close()
    return jsonify({"ok": True, "data": [dict(r) for r in rows]})


@app.route("/api/users", methods=["POST"])
def api_users_create():
    data = request.json
    conn = get_db()
    try:
        cur = conn.execute(
            "INSERT INTO users (email,password_hash,nome,cognome,ruolo,agente_id) VALUES (?,?,?,?,?,?)",
            (data["email"], hash_pw(data.get("password", "ulteria2026")),
             data["nome"], data["cognome"], data.get("ruolo", "agente"), data.get("agente_id")),
        )
        conn.commit()
        row = conn.execute("SELECT id,email,nome,cognome,ruolo,agente_id FROM users WHERE id=?", (cur.lastrowid,)).fetchone()
        conn.close()
        return jsonify({"ok": True, "data": dict(row)}), 201
    except Exception as e:
        conn.close()
        return jsonify({"ok": False, "error": str(e)}), 400


# ── Routes ───────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/oggetti/<int:oid>")
def oggetto_page(oid):
    return render_template("oggetto.html", oggetto_id=oid)


@app.route("/impostazioni")
def impostazioni_page():
    return render_template("impostazioni.html")


# ── API: Config ──────────────────────────────────────────────────────

@app.route("/api/config", methods=["GET"])
def api_config_get():
    cfg = read_config()
    conn = get_db()
    total = conn.execute("SELECT COUNT(*) FROM offerte").fetchone()[0]
    conn.close()
    cfg["totale_offerte_generate"] = total
    return jsonify(cfg)


@app.route("/api/config", methods=["POST"])
def api_config_post():
    data = request.json
    cfg = read_config()
    if "prossimo_numero" in data:
        cfg["prossimo_numero"] = int(data["prossimo_numero"])
    write_config(cfg)
    return jsonify({"ok": True})


# ── API: Offerte ─────────────────────────────────────────────────────

@app.route("/api/offerte", methods=["GET"])
def api_offerte_list():
    conn = get_db()
    sql = """SELECT o.*,
        obj.nome as oggetto_nome, obj.via as oggetto_via, obj.civico as oggetto_civico, obj.comune as oggetto_comune,
        c.nome_studio as cliente_nome, c.tipo_cliente as cliente_tipo,
        a.nome as agente_nome, a.cognome as agente_cognome, a.colore as agente_colore,
        s.nome as segnalatore_nome
        FROM offerte o
        LEFT JOIN oggetti obj ON o.oggetto_id = obj.id
        LEFT JOIN clienti c ON LOWER(o.nome_studio) = LOWER(c.nome_studio)
        LEFT JOIN agenti a ON o.agente_id = a.id
        LEFT JOIN segnalatori s ON o.segnalatore_id = s.id
        WHERE 1=1"""
    params = []
    if request.args.get("agente_id"):
        sql += " AND o.agente_id=?"
        params.append(int(request.args["agente_id"]))
    if request.args.get("stato"):
        stati = request.args["stato"].split(",")
        sql += " AND o.stato IN (" + ",".join("?" * len(stati)) + ")"
        params.extend(stati)
    if request.args.get("template"):
        sql += " AND o.template=?"
        params.append(request.args["template"])
    if request.args.get("macro_categoria"):
        sql += " AND o.macro_categoria=?"
        params.append(request.args["macro_categoria"])
    if request.args.get("sottotipo"):
        sql += " AND o.sottotipo=?"
        params.append(request.args["sottotipo"])
    if request.args.get("tipo_cliente"):
        sql += " AND c.tipo_cliente=?"
        params.append(request.args["tipo_cliente"])
    if request.args.get("q"):
        q = f"%{request.args['q']}%"
        sql += " AND (o.nome_studio LIKE ? OR o.nome_condominio LIKE ? OR o.citta LIKE ? OR CAST(o.numero AS TEXT) LIKE ?)"
        params.extend([q, q, q, q])
    if request.args.get("dal"):
        sql += " AND o.data_creazione>=?"
        params.append(request.args["dal"])
    if request.args.get("al"):
        sql += " AND o.data_creazione<=?"
        params.append(request.args["al"] + " 23:59:59")
    sql += " ORDER BY o.id DESC"
    rows = conn.execute(sql, params).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/api/offerte", methods=["POST"])
def api_offerte_create():
    data = request.json
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Assign numero progressivo
    cfg = read_config()
    numero = cfg.get("prossimo_numero", 26000)
    cfg["prossimo_numero"] = numero + 1
    write_config(cfg)

    conn = get_db()
    cur = conn.execute(
        """INSERT INTO offerte
           (numero, nome_studio, nome_condominio, via, cap, citta, riferimento,
            template, prezzo_fornitura, prezzo_care, canone_lettura,
            modalita, totale, stato, email_studio, data_creazione, note,
            agente_id, oggetto_id, macro_categoria, sottotipo,
            valore_commessa, importo, importo_servizio_annuo,
            natura, tipo_offerta, is_gara_appalto, gara_id, valore_gara,
            stato_versione, versione, segnalatore_id)
           VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            numero,
            data.get("nome_studio", ""),
            data.get("nome_condominio", ""),
            data.get("via", ""),
            data.get("cap", ""),
            data.get("citta", ""),
            data.get("riferimento", ""),
            data.get("template", ""),
            data.get("prezzo_fornitura"),
            data.get("prezzo_care"),
            data.get("canone_lettura"),
            data.get("modalita", "vendita"),
            data.get("totale"),
            data.get("stato", "richiamato"),
            data.get("email_studio", ""),
            now,
            data.get("note", ""),
            data.get("agente_id"),
            data.get("oggetto_id"),
            data.get("macro_categoria"),
            data.get("sottotipo"),
            data.get("valore_commessa"),
            data.get("importo"),
            data.get("importo_servizio_annuo"),
            data.get("natura", "nuovo"),
            data.get("tipo_offerta"),
            data.get("is_gara_appalto", 0),
            data.get("gara_id"),
            data.get("valore_gara"),
            data.get("stato_versione", "attiva"),
            data.get("versione", "A"),
            data.get("segnalatore_id"),
        ),
    )
    conn.commit()
    new_id = cur.lastrowid
    row = conn.execute("SELECT * FROM offerte WHERE id=?", (new_id,)).fetchone()
    conn.close()
    return jsonify(dict(row)), 201


@app.route("/api/offerte/<int:oid>", methods=["PUT"])
def api_offerte_update(oid):
    data = request.json
    conn = get_db()
    row = conn.execute("SELECT * FROM offerte WHERE id=?", (oid,)).fetchone()
    if not row:
        conn.close()
        return jsonify({"error": "Non trovata"}), 404

    allowed = [
        "nome_studio", "nome_condominio", "via", "cap", "citta",
        "riferimento", "template", "prezzo_fornitura", "prezzo_care",
        "canone_lettura", "modalita", "totale", "stato", "email_studio",
        "note", "path_docx", "path_pdf", "numero", "agente_id",
        "motivo_perdita", "note_perdita", "importo",
        "macro_categoria", "sottotipo", "valore_commessa",
        "is_gara_appalto", "valore_gara", "gara_id",
        "segnalatore_id", "tipo_servizio",
    ]
    sets = []
    vals = []
    for k in allowed:
        if k in data:
            sets.append(f"{k}=?")
            vals.append(data[k])
    if sets:
        vals.append(oid)
        conn.execute(f"UPDATE offerte SET {','.join(sets)} WHERE id=?", vals)
        conn.commit()

    row = conn.execute("SELECT * FROM offerte WHERE id=?", (oid,)).fetchone()
    conn.close()
    return jsonify(dict(row))


@app.route("/api/offerte/<int:oid>/stato", methods=["PUT"])
def api_offerte_stato(oid):
    data = request.json
    conn = get_db()
    conn.execute("UPDATE offerte SET stato=? WHERE id=?", (data["stato"], oid))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/offerte/<int:oid>", methods=["DELETE"])
def api_offerte_delete(oid):
    conn = get_db()
    conn.execute("DELETE FROM offerte WHERE id=?", (oid,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/offerte/<int:oid>/duplica", methods=["POST"])
def api_offerte_duplica(oid):
    conn = get_db()
    row = conn.execute("SELECT * FROM offerte WHERE id=?", (oid,)).fetchone()
    if not row:
        conn.close()
        return jsonify({"error": "Non trovata"}), 404
    orig = dict(row)

    cfg = read_config()
    numero = cfg["prossimo_numero"]
    cfg["prossimo_numero"] = numero + 1
    write_config(cfg)

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cur = conn.execute(
        """INSERT INTO offerte
           (numero, nome_studio, nome_condominio, via, cap, citta,
            riferimento, template, prezzo_fornitura, prezzo_care,
            canone_lettura, modalita, importo, stato, email_studio,
            data_creazione, note, agente_id)
           VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            numero, orig["nome_studio"], orig["nome_condominio"],
            orig["via"], orig["cap"], orig["citta"],
            orig["riferimento"], orig["template"],
            orig["prezzo_fornitura"], orig["prezzo_care"],
            orig["canone_lettura"], orig["modalita"], orig.get("importo"),
            "in_attesa_assemblea", orig["email_studio"],
            now, orig["note"], orig["agente_id"],
        ),
    )
    conn.commit()
    new_row = conn.execute("SELECT * FROM offerte WHERE id=?", (cur.lastrowid,)).fetchone()
    conn.close()
    return jsonify(dict(new_row)), 201


# ── API: Genera ──────────────────────────────────────────────────────

@app.route("/api/genera", methods=["POST"])
def api_genera():
    data = request.json
    oid = data.get("id")
    conn = get_db()
    row = conn.execute("SELECT * FROM offerte WHERE id=?", (oid,)).fetchone()
    if not row:
        conn.close()
        return jsonify({"error": "Offerta non trovata"}), 404
    off = dict(row)

    # Validate required fields
    missing = []
    for f in ["nome_studio", "template", "prezzo_fornitura", "prezzo_care", "canone_lettura"]:
        if not off.get(f) and off.get(f) != 0:
            missing.append(f)
    if missing:
        conn.close()
        return jsonify({"error": f"Campi obbligatori mancanti: {', '.join(missing)}"}), 400

    template_key = off["template"]
    if template_key not in TEMPLATE_MAP:
        conn.close()
        return jsonify({"error": f"Template sconosciuto: {template_key}"}), 400

    # Assign number
    cfg = read_config()
    if not off.get("numero"):
        numero = cfg["prossimo_numero"]
        cfg["prossimo_numero"] = numero + 1
        write_config(cfg)
    else:
        numero = off["numero"]

    slug = studio_slug(off["nome_studio"])
    tipo = template_key
    anno = datetime.now().strftime("%Y")
    folder_name = f"{numero}_{slug}_{tipo}"
    dest_dir = OUTPUT_DIR / anno / folder_name
    dest_dir.mkdir(parents=True, exist_ok=True)

    file_base = f"ULTERIA_{numero}_{slug}_{tipo}"
    docx_name = f"{file_base}.docx"
    docx_path = dest_dir / docx_name

    # Copy template
    src = UPLOADS_DIR / TEMPLATE_MAP[template_key]
    shutil.copy2(str(src), str(docx_path))

    # Open and replace
    doc = Document(str(docx_path))

    via_studio = off.get("via", "") or ""

    dat_str = datetime.now().strftime("%d/%m/%Y")

    pfo = format_eur(off.get("prezzo_fornitura"))
    pca = format_eur(off.get("prezzo_care"))
    pcl = format_eur(off.get("canone_lettura"))
    modalita = off.get("modalita", "vendita")
    if modalita == "comodato":
        mod_val = "COMODATO D'USO"
    else:
        totale = (off.get("prezzo_fornitura") or 0) + (off.get("prezzo_care") or 0) + (off.get("canone_lettura") or 0)
        mod_val = f"\u20ac {format_eur(totale)}"

    # Body paragraphs
    replace_runs(doc.paragraphs, "\u00abSTU\u00bb", off["nome_studio"])
    replace_runs(doc.paragraphs, "\u00abVIA\u00bb", via_studio)
    replace_runs(doc.paragraphs, "\u00abDAT\u00bb", dat_str)

    # Fix STU/VIA alignment: put VIA on new line under STU
    fix_stu_via_paragraph(doc, off["nome_studio"], via_studio)

    # All sections (headers, footers, textboxes) via XML
    for section in doc.sections:
        for rel_type in [section.header, section.footer,
                         section.first_page_header, section.first_page_footer,
                         section.even_page_header, section.even_page_footer]:
            if rel_type and rel_type.is_linked_to_previous is False or rel_type:
                try:
                    replace_xml(rel_type._element, "\u00abNR\u00bb", str(numero))
                    replace_xml(rel_type._element, "\u00abSTU\u00bb", off["nome_studio"])
                    replace_xml(rel_type._element, "\u00abVIA\u00bb", via_studio)
                    replace_xml(rel_type._element, "\u00abDAT\u00bb", dat_str)
                except Exception:
                    pass

    # Textboxes in body XML
    replace_xml(doc.element.body, "\u00abPFO\u00bb", pfo)
    replace_xml(doc.element.body, "\u00abPCA\u00bb", pca)
    replace_xml(doc.element.body, "\u00abPCL\u00bb", pcl)
    replace_xml(doc.element.body, "\u00abMOD\u00bb", mod_val)
    replace_xml(doc.element.body, "\u00abNR\u00bb", str(numero))
    replace_xml(doc.element.body, "\u00abSTU\u00bb", off["nome_studio"])
    replace_xml(doc.element.body, "\u00abVIA\u00bb", via_studio)
    replace_xml(doc.element.body, "\u00abDAT\u00bb", dat_str)

    doc.save(str(docx_path))

    # PDF conversion via Word COM
    pdf_error = False
    pdf_path = dest_dir / f"{file_base}.pdf"
    pdf_rel = None
    try:
        import pythoncom
        pythoncom.CoInitialize()
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            doc_com = word.Documents.Open(str(docx_path.resolve()))
            doc_com.SaveAs2(str(pdf_path.resolve()), FileFormat=17)  # 17 = PDF
            doc_com.Close(False)
            word.Quit()
            pdf_rel = f"/output/{anno}/{folder_name}/{file_base}.pdf"
        finally:
            pythoncom.CoUninitialize()
    except Exception as e:
        pdf_error = True
        app.logger.warning(f"PDF conversion failed: {e}")
        # Fallback: try docx2pdf
        try:
            from docx2pdf import convert
            convert(str(docx_path), str(pdf_path))
            pdf_rel = f"/output/{anno}/{folder_name}/{file_base}.pdf"
            pdf_error = False
        except Exception as e2:
            app.logger.warning(f"docx2pdf fallback also failed: {e2}")

    docx_rel = f"/output/{anno}/{folder_name}/{docx_name}"

    # Update DB
    conn.execute(
        "UPDATE offerte SET numero=?, path_docx=?, path_pdf=? WHERE id=?",
        (numero, docx_rel, pdf_rel, oid),
    )
    conn.commit()
    conn.close()

    return jsonify({
        "ok": True,
        "numero": numero,
        "nome_file": file_base,
        "docx_url": docx_rel,
        "pdf_url": pdf_rel,
        "pdf_error": pdf_error,
    })


# ── Serve output files ───────────────────────────────────────────────

@app.route("/output/<path:filepath>")
def serve_output(filepath):
    full = OUTPUT_DIR / filepath
    if not full.exists():
        abort(404)
    directory = str(full.parent)
    filename = full.name
    return send_from_directory(directory, filename)


# ── API: Clienti ─────────────────────────────────────────────────────

@app.route("/api/clienti", methods=["GET"])
def api_clienti_list():
    q = request.args.get("q", "").strip()
    conn = get_db()
    if q and len(q) >= 2:
        rows = conn.execute(
            "SELECT * FROM clienti WHERE nome_studio LIKE ? OR citta LIKE ? OR referente LIKE ? ORDER BY nome_studio",
            (f"%{q}%", f"%{q}%", f"%{q}%"),
        ).fetchall()
    else:
        rows = conn.execute("SELECT * FROM clienti ORDER BY nome_studio").fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/api/clienti", methods=["POST"])
def api_clienti_create():
    data = request.json
    nome = data.get("nome_studio", "").strip()
    if not nome:
        return jsonify({"error": "nome_studio obbligatorio"}), 400

    conn = get_db()
    # Check if client already exists (case-insensitive)
    existing = conn.execute(
        "SELECT * FROM clienti WHERE LOWER(nome_studio) = LOWER(?)", (nome,)
    ).fetchone()
    if existing:
        conn.close()
        return jsonify(dict(existing)), 200  # Return existing, don't duplicate

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cur = conn.execute(
        """INSERT INTO clienti
           (nome_studio, via, cap, citta, email, telefono, referente, note, data_inserimento)
           VALUES (?,?,?,?,?,?,?,?,?)""",
        (
            nome,
            data.get("via", ""),
            data.get("cap", ""),
            data.get("citta", ""),
            data.get("email", ""),
            data.get("telefono", ""),
            data.get("referente", ""),
            data.get("note", ""),
            now,
        ),
    )
    conn.commit()
    new_id = cur.lastrowid
    row = conn.execute("SELECT * FROM clienti WHERE id=?", (new_id,)).fetchone()
    conn.close()
    return jsonify(dict(row)), 201


@app.route("/api/clienti/<int:cid>", methods=["GET"])
def api_clienti_detail(cid):
    conn = get_db()
    cliente = conn.execute("SELECT * FROM clienti WHERE id=?", (cid,)).fetchone()
    if not cliente:
        conn.close()
        return jsonify({"error": "Cliente non trovato"}), 404
    offerte = conn.execute(
        "SELECT * FROM offerte WHERE nome_studio=? ORDER BY data_creazione DESC",
        (cliente["nome_studio"],),
    ).fetchall()
    conn.close()
    return jsonify({
        "cliente": dict(cliente),
        "offerte": [dict(o) for o in offerte],
    })


@app.route("/api/clienti/<int:cid>", methods=["PUT"])
def api_clienti_update(cid):
    data = request.json
    conn = get_db()
    allowed = ["nome_studio", "via", "cap", "citta", "email", "telefono", "referente", "note", "tipo_cliente"]
    sets = []
    vals = []
    for k in allowed:
        if k in data:
            sets.append(f"{k}=?")
            vals.append(data[k])
    if sets:
        vals.append(cid)
        conn.execute(f"UPDATE clienti SET {','.join(sets)} WHERE id=?", vals)
        conn.commit()
    row = conn.execute("SELECT * FROM clienti WHERE id=?", (cid,)).fetchone()
    conn.close()
    return jsonify(dict(row))


# ── Scheda Cliente (pagina dedicata) ─────────────────────────────────

@app.route("/clienti/<int:cid>")
def cliente_page(cid):
    return render_template("cliente.html", cliente_id=cid)


@app.route("/api/clienti/<int:cid>/note", methods=["POST"])
def api_clienti_add_note(cid):
    data = request.json
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn = get_db()
    cur = conn.execute(
        "INSERT INTO note_clienti (cliente_id, testo, autore, created_at) VALUES (?,?,?,?)",
        (cid, data.get("testo", ""), data.get("autore", ""), now),
    )
    conn.commit()
    row = conn.execute("SELECT * FROM note_clienti WHERE id=?", (cur.lastrowid,)).fetchone()
    conn.close()
    return jsonify({"ok": True, "data": dict(row)}), 201


@app.route("/api/note/<int:nid>", methods=["DELETE"])
def api_note_delete(nid):
    conn = get_db()
    conn.execute("DELETE FROM note_clienti WHERE id=?", (nid,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/clienti/<int:cid>/full", methods=["GET"])
def api_clienti_full(cid):
    """Full client detail with offerte, condomini, note."""
    conn = get_db()
    cliente = conn.execute("SELECT * FROM clienti WHERE id=?", (cid,)).fetchone()
    if not cliente:
        conn.close()
        return jsonify({"error": "Cliente non trovato"}), 404
    offerte = conn.execute(
        "SELECT o.*, a.nome as agente_nome, a.cognome as agente_cognome, a.colore as agente_colore "
        "FROM offerte o LEFT JOIN agenti a ON o.agente_id = a.id "
        "WHERE LOWER(o.nome_studio) = LOWER(?) ORDER BY o.data_creazione DESC",
        (cliente["nome_studio"],),
    ).fetchall()
    condomini = conn.execute(
        "SELECT * FROM condomini WHERE cliente_id=? ORDER BY nome", (cid,)
    ).fetchall()
    note = conn.execute(
        "SELECT * FROM note_clienti WHERE cliente_id=? ORDER BY created_at DESC", (cid,)
    ).fetchall()
    conn.close()
    return jsonify({
        "cliente": dict(cliente),
        "offerte": [dict(o) for o in offerte],
        "condomini": [dict(c) for c in condomini],
        "note": [dict(n) for n in note],
    })


# ── API: Agenti ──────────────────────────────────────────────────────

@app.route("/api/agenti", methods=["GET"])
def api_agenti_list():
    conn = get_db()
    rows = conn.execute("SELECT * FROM agenti ORDER BY cognome, nome").fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/api/agenti", methods=["POST"])
def api_agenti_create():
    data = request.json
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn = get_db()
    cur = conn.execute(
        """INSERT INTO agenti (nome, cognome, email, telefono, colore, note, data_inserimento)
           VALUES (?,?,?,?,?,?,?)""",
        (
            data.get("nome", ""),
            data.get("cognome", ""),
            data.get("email", ""),
            data.get("telefono", ""),
            data.get("colore", "#009FE3"),
            data.get("note", ""),
            now,
        ),
    )
    conn.commit()
    row = conn.execute("SELECT * FROM agenti WHERE id=?", (cur.lastrowid,)).fetchone()
    conn.close()
    return jsonify(dict(row)), 201


@app.route("/api/agenti/<int:aid>", methods=["GET"])
def api_agenti_detail(aid):
    conn = get_db()
    agente = conn.execute("SELECT * FROM agenti WHERE id=?", (aid,)).fetchone()
    if not agente:
        conn.close()
        return jsonify({"error": "Agente non trovato"}), 404
    offerte = conn.execute(
        "SELECT * FROM offerte WHERE agente_id=? ORDER BY data_creazione DESC",
        (aid,),
    ).fetchall()
    conn.close()
    return jsonify({
        "agente": dict(agente),
        "offerte": [dict(o) for o in offerte],
    })


@app.route("/api/agenti/<int:aid>", methods=["PUT"])
def api_agenti_update(aid):
    data = request.json
    conn = get_db()
    allowed = ["nome", "cognome", "email", "telefono", "colore", "note"]
    sets = []
    vals = []
    for k in allowed:
        if k in data:
            sets.append(f"{k}=?")
            vals.append(data[k])
    if sets:
        vals.append(aid)
        conn.execute(f"UPDATE agenti SET {','.join(sets)} WHERE id=?", vals)
        conn.commit()
    row = conn.execute("SELECT * FROM agenti WHERE id=?", (aid,)).fetchone()
    conn.close()
    return jsonify(dict(row))


# ── API: Agenti Dashboard ────────────────────────────────────────────

@app.route("/agenti/<int:aid>")
def agente_page(aid):
    return render_template("agente.html", agente_id=aid)


@app.route("/api/agenti/<int:aid>/dashboard", methods=["GET"])
def api_agente_dashboard(aid):
    conn = get_db()
    agente = conn.execute("SELECT * FROM agenti WHERE id=?", (aid,)).fetchone()
    if not agente:
        conn.close()
        return jsonify({"error": "Agente non trovato"}), 404

    offerte = conn.execute(
        "SELECT * FROM offerte WHERE agente_id=? ORDER BY data_creazione DESC", (aid,)
    ).fetchall()

    attivita = conn.execute(
        "SELECT a.*, c.nome_studio as cliente_nome FROM attivita a "
        "LEFT JOIN clienti c ON a.cliente_id = c.id "
        "WHERE a.agente_id=? ORDER BY a.data_scadenza ASC", (aid,)
    ).fetchall()

    # Clienti affidati: clienti con almeno un'offerta di questo agente
    clienti_rows = conn.execute(
        "SELECT c.*, COUNT(o.id) as num_offerte, MAX(o.data_creazione) as ultimo_contatto "
        "FROM clienti c JOIN offerte o ON LOWER(c.nome_studio) = LOWER(o.nome_studio) "
        "WHERE o.agente_id=? GROUP BY c.id ORDER BY c.nome_studio", (aid,)
    ).fetchall()

    # Stats mensili (ultimi 6 mesi)
    stats = conn.execute(
        "SELECT strftime('%Y-%m', data_creazione) as mese, "
        "COUNT(*) as inviate, "
        "SUM(CASE WHEN stato='preso_lavoro' THEN 1 ELSE 0 END) as prese, "
        "SUM(CASE WHEN stato='perso' THEN 1 ELSE 0 END) as perse "
        "FROM offerte WHERE agente_id=? AND data_creazione IS NOT NULL "
        "GROUP BY mese ORDER BY mese DESC LIMIT 6", (aid,)
    ).fetchall()

    conn.close()
    return jsonify({
        "agente": dict(agente),
        "offerte": [dict(o) for o in offerte],
        "attivita": [dict(a) for a in attivita],
        "clienti": [dict(c) for c in clienti_rows],
        "stats": [dict(s) for s in stats],
    })


# ── API: Attivita ────────────────────────────────────────────────────

@app.route("/api/attivita", methods=["POST"])
def api_attivita_create():
    data = request.json
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn = get_db()
    cur = conn.execute(
        """INSERT INTO attivita (agente_id, tipo, titolo, descrizione,
           data_scadenza, priorita, stato, cliente_id, offerta_id, created_at)
           VALUES (?,?,?,?,?,?,?,?,?,?)""",
        (
            data.get("agente_id"),
            data.get("tipo", "todo"),
            data.get("titolo", ""),
            data.get("descrizione", ""),
            data.get("data_scadenza"),
            data.get("priorita", "media"),
            "aperta",
            data.get("cliente_id"),
            data.get("offerta_id"),
            now,
        ),
    )
    conn.commit()
    row = conn.execute("SELECT * FROM attivita WHERE id=?", (cur.lastrowid,)).fetchone()
    conn.close()
    return jsonify(dict(row)), 201


@app.route("/api/attivita/<int:tid>", methods=["PATCH"])
def api_attivita_update(tid):
    data = request.json
    conn = get_db()
    allowed = ["tipo", "titolo", "descrizione", "data_scadenza", "priorita", "cliente_id", "offerta_id"]
    sets = []
    vals = []
    for k in allowed:
        if k in data:
            sets.append(f"{k}=?")
            vals.append(data[k])
    if sets:
        vals.append(tid)
        conn.execute(f"UPDATE attivita SET {','.join(sets)} WHERE id=?", vals)
        conn.commit()
    row = conn.execute("SELECT * FROM attivita WHERE id=?", (tid,)).fetchone()
    conn.close()
    return jsonify(dict(row))


@app.route("/api/attivita/<int:tid>/completa", methods=["PATCH"])
def api_attivita_completa(tid):
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn = get_db()
    conn.execute(
        "UPDATE attivita SET stato='completata', completato_il=? WHERE id=?",
        (now, tid),
    )
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/attivita/<int:tid>", methods=["DELETE"])
def api_attivita_delete(tid):
    conn = get_db()
    conn.execute("DELETE FROM attivita WHERE id=?", (tid,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/clienti/search", methods=["GET"])
def api_clienti_search():
    q = request.args.get("q", "").strip()
    conn = get_db()
    if q and len(q) >= 1:
        rows = conn.execute(
            "SELECT id, nome_studio, citta FROM clienti WHERE nome_studio LIKE ? ORDER BY nome_studio LIMIT 10",
            (f"%{q}%",),
        ).fetchall()
    else:
        rows = conn.execute("SELECT id, nome_studio, citta FROM clienti ORDER BY nome_studio LIMIT 10").fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/api/offerte/by-cliente/<int:cid>", methods=["GET"])
def api_offerte_by_cliente(cid):
    conn = get_db()
    cliente = conn.execute("SELECT nome_studio FROM clienti WHERE id=?", (cid,)).fetchone()
    if not cliente:
        conn.close()
        return jsonify([])
    rows = conn.execute(
        "SELECT id, numero, nome_condominio, riferimento FROM offerte WHERE LOWER(nome_studio)=LOWER(?) ORDER BY data_creazione DESC",
        (cliente["nome_studio"],),
    ).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/api/agenti/<int:aid>/badges", methods=["GET"])
def api_agente_badges(aid):
    conn = get_db()
    offerte_aperte = conn.execute(
        "SELECT COUNT(*) FROM offerte WHERE agente_id=? AND stato IN ('richiamato','in_attesa_assemblea','rimandato')",
        (aid,),
    ).fetchone()[0]
    attivita_urgenti = conn.execute(
        "SELECT COUNT(*) FROM attivita WHERE agente_id=? AND stato='aperta' AND data_scadenza <= date('now')",
        (aid,),
    ).fetchone()[0]
    conn.close()
    return jsonify({"offerte_aperte": offerte_aperte, "attivita_urgenti": attivita_urgenti})


@app.route("/api/attivita/scadute-count", methods=["GET"])
def api_attivita_scadute_count():
    conn = get_db()
    count = conn.execute(
        "SELECT COUNT(*) FROM attivita WHERE stato='aperta' AND date(data_scadenza) <= date('now')"
    ).fetchone()[0]
    conn.close()
    return jsonify({"count": count})


@app.route("/api/agenti/<int:aid>/stats", methods=["GET"])
def api_agente_stats(aid):
    conn = get_db()
    offs = conn.execute("SELECT * FROM offerte WHERE agente_id=?", (aid,)).fetchall()
    tot = len(offs)
    prese = sum(1 for o in offs if o["stato"] == "preso_lavoro")
    perse = sum(1 for o in offs if o["stato"] == "perso")
    aperte = sum(1 for o in offs if o["stato"] in ("richiamato", "in_attesa_assemblea", "rimandato"))
    val_preso = sum((o["prezzo_fornitura"] or 0) + (o["prezzo_care"] or 0) + (o["canone_lettura"] or 0) for o in offs if o["stato"] == "preso_lavoro")
    val_pipeline = sum((o["prezzo_fornitura"] or 0) + (o["prezzo_care"] or 0) + (o["canone_lettura"] or 0) for o in offs if o["stato"] in ("richiamato", "in_attesa_assemblea"))
    tasso = round(prese / tot * 100, 1) if tot > 0 else 0
    att_aperte = conn.execute("SELECT COUNT(*) FROM attivita WHERE agente_id=? AND stato='aperta'", (aid,)).fetchone()[0]
    att_scadute = conn.execute("SELECT COUNT(*) FROM attivita WHERE agente_id=? AND stato='aperta' AND data_scadenza <= date('now')", (aid,)).fetchone()[0]
    conn.close()
    return jsonify({
        "ok": True,
        "data": {
            "totale_offerte": tot, "prese": prese, "perse": perse, "aperte": aperte,
            "valore_preso": val_preso, "valore_pipeline": val_pipeline, "tasso_conversione": tasso,
            "attivita_aperte": att_aperte, "attivita_scadute": att_scadute,
        }
    })


# ── Step 8: API Etichette ────────────────────────────────────────────

@app.route("/api/etichette", methods=["GET"])
def api_etichette_list():
    cat = request.args.get("categoria", "")
    conn = get_db()
    if cat:
        rows = conn.execute(
            "SELECT * FROM etichette WHERE categoria=? AND attiva=1 ORDER BY ordine, valore", (cat,)
        ).fetchall()
    else:
        rows = conn.execute("SELECT * FROM etichette WHERE attiva=1 ORDER BY categoria, ordine, valore").fetchall()
    conn.close()
    return jsonify({"ok": True, "data": [dict(r) for r in rows]})


@app.route("/api/etichette", methods=["POST"])
def api_etichette_create():
    data = request.json
    conn = get_db()
    cur = conn.execute(
        "INSERT INTO etichette (categoria, valore, colore_bg, colore_testo, ordine) VALUES (?,?,?,?,?)",
        (data.get("categoria"), data.get("valore"), data.get("colore_bg", "#F4F9FD"),
         data.get("colore_testo", "#0D1F35"), data.get("ordine", 0)),
    )
    conn.commit()
    row = conn.execute("SELECT * FROM etichette WHERE id=?", (cur.lastrowid,)).fetchone()
    conn.close()
    return jsonify({"ok": True, "data": dict(row)}), 201


@app.route("/api/etichette/<int:eid>", methods=["PATCH"])
def api_etichette_update(eid):
    data = request.json
    conn = get_db()
    allowed = ["valore", "colore_bg", "colore_testo", "ordine", "attiva"]
    sets, vals = [], []
    for k in allowed:
        if k in data:
            sets.append(f"{k}=?")
            vals.append(data[k])
    if sets:
        vals.append(eid)
        conn.execute(f"UPDATE etichette SET {','.join(sets)} WHERE id=?", vals)
        conn.commit()
    row = conn.execute("SELECT * FROM etichette WHERE id=?", (eid,)).fetchone()
    conn.close()
    return jsonify({"ok": True, "data": dict(row)})


@app.route("/api/etichette/<int:eid>", methods=["DELETE"])
def api_etichette_delete(eid):
    conn = get_db()
    conn.execute("DELETE FROM etichette WHERE id=?", (eid,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


# ── API: Oggetti ─────────────────────────────────────────────────────

@app.route("/api/oggetti", methods=["POST"])
def api_oggetti_create():
    data = request.json
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn = get_db()
    cur = conn.execute(
        """INSERT INTO oggetti (cliente_id, nome, via, civico, comune, provincia, cap,
           tipo_oggetto, stato_pipeline, agente_id, natura, n_unita, n_scale, created_at)
           VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (data.get("cliente_id"), data.get("nome"), data.get("via"), data.get("civico"),
         data.get("comune"), data.get("provincia"), data.get("cap"),
         data.get("tipo_oggetto", "condominio"), data.get("stato_pipeline", "prospect"),
         data.get("agente_id"), data.get("natura"), data.get("n_unita"), data.get("n_scale"), now),
    )
    conn.commit()
    row = conn.execute("SELECT * FROM oggetti WHERE id=?", (cur.lastrowid,)).fetchone()
    conn.close()
    return jsonify({"ok": True, "data": dict(row)}), 201


@app.route("/api/oggetti/<int:oid>", methods=["GET"])
def api_oggetti_detail(oid):
    conn = get_db()
    obj = conn.execute("SELECT * FROM oggetti WHERE id=?", (oid,)).fetchone()
    if not obj:
        conn.close()
        return jsonify({"ok": False, "error": "Oggetto non trovato"}), 404
    cliente = conn.execute("SELECT * FROM clienti WHERE id=?", (obj["cliente_id"],)).fetchone()
    offerte = conn.execute(
        "SELECT o.*, a.nome as agente_nome, a.cognome as agente_cognome, a.colore as agente_colore "
        "FROM offerte o LEFT JOIN agenti a ON o.agente_id = a.id "
        "WHERE o.oggetto_id=? ORDER BY o.data_creazione DESC", (oid,)
    ).fetchall()
    note = conn.execute(
        "SELECT * FROM note_clienti WHERE oggetto_id=? ORDER BY created_at DESC", (oid,)
    ).fetchall()
    attivita = conn.execute(
        "SELECT a.*, c.nome_studio as cliente_nome FROM attivita a "
        "LEFT JOIN clienti c ON a.cliente_id = c.id "
        "WHERE a.oggetto_id=? ORDER BY a.data_scadenza ASC", (oid,)
    ).fetchall()
    timeline = conn.execute(
        "SELECT * FROM timeline_eventi WHERE oggetto_id=? ORDER BY created_at DESC LIMIT 50", (oid,)
    ).fetchall()
    conn.close()
    return jsonify({
        "ok": True,
        "data": {
            "oggetto": dict(obj),
            "cliente": dict(cliente) if cliente else None,
            "offerte": [dict(o) for o in offerte],
            "note": [dict(n) for n in note],
            "attivita": [dict(a) for a in attivita],
            "timeline": [dict(t) for t in timeline],
        }
    })


@app.route("/api/oggetti/<int:oid>", methods=["PATCH"])
def api_oggetti_update(oid):
    data = request.json
    conn = get_db()
    allowed = ["nome", "via", "civico", "comune", "provincia", "cap",
               "tipo_oggetto", "stato_pipeline", "agente_id", "natura",
               "n_unita", "n_scale", "motivo_perdita", "note_perdita",
               "data_primo_contatto", "data_ultimo_contatto"]
    sets, vals = [], []
    for k in allowed:
        if k in data:
            sets.append(f"{k}=?")
            vals.append(data[k])
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    sets.append("updated_at=?")
    vals.append(now)
    vals.append(oid)
    conn.execute(f"UPDATE oggetti SET {','.join(sets)} WHERE id=?", vals)
    conn.commit()
    row = conn.execute("SELECT * FROM oggetti WHERE id=?", (oid,)).fetchone()
    conn.close()
    return jsonify({"ok": True, "data": dict(row)})


@app.route("/api/oggetti/<int:oid>/note", methods=["POST"])
def api_oggetti_add_note(oid):
    data = request.json
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn = get_db()
    cur = conn.execute(
        "INSERT INTO note_clienti (oggetto_id, testo, autore, created_at) VALUES (?,?,?,?)",
        (oid, data.get("testo", ""), data.get("autore", ""), now),
    )
    conn.commit()
    row = conn.execute("SELECT * FROM note_clienti WHERE id=?", (cur.lastrowid,)).fetchone()
    conn.close()
    return jsonify({"ok": True, "data": dict(row)}), 201


@app.route("/api/oggetti/<int:oid>/intestazione", methods=["POST"])
def api_oggetti_intestazione(oid):
    data = request.json
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn = get_db()
    obj = conn.execute("SELECT * FROM oggetti WHERE id=?", (oid,)).fetchone()
    if not obj:
        conn.close()
        return jsonify({"ok": False, "error": "Non trovato"}), 404
    old_cliente_id = obj["cliente_id"]
    new_cliente_id = data.get("nuovo_cliente_id")
    conn.execute(
        "UPDATE oggetti SET cliente_id=?, cliente_precedente_id=?, data_cambio_intestazione=?, updated_at=? WHERE id=?",
        (new_cliente_id, old_cliente_id, data.get("data_cambio", now[:10]), now, oid),
    )
    conn.execute(
        "INSERT INTO timeline_eventi (tipo_evento, descrizione, oggetto_id, utente, created_at) VALUES (?,?,?,?,?)",
        ("intestazione_cambiata", "Intestazione cambiata", oid, data.get("utente", ""), now),
    )
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/oggetti/by-cliente/<int:cid>", methods=["GET"])
def api_oggetti_by_cliente(cid):
    conn = get_db()
    rows = conn.execute("SELECT * FROM oggetti WHERE cliente_id=? ORDER BY comune, via", (cid,)).fetchall()
    conn.close()
    return jsonify({"ok": True, "data": [dict(r) for r in rows]})


# ── Offerte versioning ───────────────────────────────────────────────

@app.route("/api/offerte/<int:oid>/versione", methods=["POST"])
def api_offerte_versione(oid):
    conn = get_db()
    orig = conn.execute("SELECT * FROM offerte WHERE id=?", (oid,)).fetchone()
    if not orig:
        conn.close()
        return jsonify({"ok": False, "error": "Non trovata"}), 404
    orig = dict(orig)

    base_id = orig.get("offerta_padre_id") or oid
    count = conn.execute(
        "SELECT COUNT(*) FROM offerte WHERE offerta_padre_id=? OR id=?", (base_id, base_id)
    ).fetchone()[0]
    next_ver = chr(ord("A") + count)

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cur = conn.execute(
        """INSERT INTO offerte
           (numero, nome_studio, nome_condominio, via, cap, citta, riferimento, template,
            prezzo_fornitura, prezzo_care, canone_lettura, modalita, importo, importo_servizio_annuo,
            stato, email_studio, data_creazione, note, agente_id, oggetto_id,
            versione, offerta_padre_id, stato_versione, tipo_offerta, natura, is_accordo_quadro)
           VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (orig["numero"], orig["nome_studio"], orig["nome_condominio"],
         orig["via"], orig["cap"], orig["citta"], orig["riferimento"], orig["template"],
         orig["prezzo_fornitura"], orig["prezzo_care"], orig["canone_lettura"],
         orig["modalita"], orig.get("importo"), orig.get("importo_servizio_annuo"),
         orig["stato"], orig["email_studio"], now, orig["note"], orig["agente_id"],
         orig.get("oggetto_id"), next_ver, base_id, "attiva",
         orig.get("tipo_offerta", "installazione"), orig.get("natura", "nuovo"),
         orig.get("is_accordo_quadro", 0)),
    )
    conn.execute("UPDATE offerte SET stato_versione='aggiornata' WHERE id=?", (oid,))
    conn.execute(
        "INSERT INTO timeline_eventi (tipo_evento, descrizione, offerta_id, oggetto_id, utente, created_at) VALUES (?,?,?,?,?,?)",
        ("offerta_aggiornata", "Versione " + next_ver + " creata", cur.lastrowid, orig.get("oggetto_id"), "", now),
    )
    conn.commit()
    row = conn.execute("SELECT * FROM offerte WHERE id=?", (cur.lastrowid,)).fetchone()
    conn.close()
    return jsonify({"ok": True, "data": dict(row)}), 201


# ── Dashboard Admin API ──────────────────────────────────────────────

@app.route("/api/dashboard/admin", methods=["GET"])
def api_dashboard_admin():
    conn = get_db()
    year = datetime.now().strftime("%Y")
    offs = conn.execute(
        "SELECT * FROM offerte WHERE stato_versione='attiva' AND strftime('%Y', data_creazione)=?", (year,)
    ).fetchall()

    total = len(offs)
    aperte = sum(1 for o in offs if o["stato"] in ("richiamato", "in_attesa_assemblea", "rimandato"))
    prese = sum(1 for o in offs if o["stato"] == "preso_lavoro")
    val_fornitura = sum((o["importo"] or 0) for o in offs if o["stato"] == "preso_lavoro" and o["tipo_offerta"] in ("fornitura", "installazione") and not o["is_accordo_quadro"])
    val_servizi = sum((o["importo_servizio_annuo"] or 0) for o in offs if o["stato"] == "preso_lavoro" and o["tipo_offerta"] == "servizio")
    aq_count = sum(1 for o in offs if o["is_accordo_quadro"])
    tasso = round(prese / total * 100, 1) if total > 0 else 0

    # Per natura
    natura_stats = {}
    for o in offs:
        n = o["natura"] or "nuovo"
        if n not in natura_stats:
            natura_stats[n] = {"count": 0, "valore": 0}
        natura_stats[n]["count"] += 1
        natura_stats[n]["valore"] += (o["importo"] or 0)

    # Per agente
    agenti_rows = conn.execute("SELECT * FROM agenti ORDER BY cognome").fetchall()
    agenti_perf = []
    for a in agenti_rows:
        a_offs = [o for o in offs if o["agente_id"] == a["id"]]
        a_prese = sum(1 for o in a_offs if o["stato"] == "preso_lavoro")
        a_val = sum((o["importo"] or 0) for o in a_offs if o["stato"] == "preso_lavoro")
        a_prospect = sum((o["importo"] or 0) for o in a_offs if o["stato"] in ("richiamato", "in_attesa_assemblea"))
        a_att = conn.execute("SELECT COUNT(*) FROM attivita WHERE agente_id=? AND stato='aperta'", (a["id"],)).fetchone()[0]
        agenti_perf.append({
            "id": a["id"], "nome": a["nome"], "cognome": a["cognome"], "colore": a["colore"],
            "offerte_mese": sum(1 for o in a_offs if o["data_creazione"] and o["data_creazione"][:7] == datetime.now().strftime("%Y-%m")),
            "offerte_ytd": len(a_offs), "valore_preso": a_val, "valore_prospect": a_prospect,
            "tasso": round(a_prese / len(a_offs) * 100, 1) if a_offs else 0,
            "attivita_aperte": a_att,
        })

    conn.close()
    return jsonify({
        "ok": True,
        "data": {
            "totale": total, "aperte": aperte, "prese": prese,
            "val_fornitura": val_fornitura, "val_servizi": val_servizi,
            "aq_count": aq_count, "tasso": tasso,
            "natura": natura_stats, "agenti": agenti_perf,
        }
    })


# ── Prompt 1: Prodotti page route ────────────────────────────────────

@app.route("/prodotti")
def prodotti_page():
    return render_template("prodotti.html")


# ── Prompt 1: API Prodotti ───────────────────────────────────────────

@app.route("/api/prodotti", methods=["GET"])
def api_prodotti_list():
    conn = get_db()
    sql = "SELECT * FROM prodotti WHERE 1=1"
    params = []
    cat = request.args.get("categoria")
    if cat:
        sql += " AND categoria=?"
        params.append(cat)
    q = request.args.get("q", "").strip()
    if q:
        sql += " AND (nome LIKE ? OR codice LIKE ?)"
        params.extend(["%" + q + "%", "%" + q + "%"])
    if request.args.get("da_aggiornare"):
        sql += " AND (data_ultimo_prezzo IS NULL OR data_ultimo_prezzo < date('now','-6 months'))"
    if not request.args.get("includi_inattivi"):
        sql += " AND attivo=1"
    sql += " ORDER BY categoria, ordine, nome"
    rows = conn.execute(sql, params).fetchall()
    conn.close()
    return jsonify({"ok": True, "data": [dict(r) for r in rows]})


@app.route("/api/prodotti", methods=["POST"])
def api_prodotti_create():
    data = request.json
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn = get_db()
    try:
        cur = conn.execute(
            """INSERT INTO prodotti (codice,nome,categoria,modello,trasmissione,dn,
               prezzo_acquisto,prezzo_vendita,data_ultimo_prezzo,fornitore,note,created_at,updated_at)
               VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (data.get("codice", "").upper(), data.get("nome"), data.get("categoria"),
             data.get("modello"), data.get("trasmissione"), data.get("dn"),
             data.get("prezzo_acquisto"), data.get("prezzo_vendita"),
             data.get("data_ultimo_prezzo", now[:10]), data.get("fornitore"),
             data.get("note"), now, now),
        )
        conn.commit()
        row = conn.execute("SELECT * FROM prodotti WHERE id=?", (cur.lastrowid,)).fetchone()
        conn.close()
        return jsonify({"ok": True, "data": dict(row)}), 201
    except Exception as e:
        conn.close()
        return jsonify({"ok": False, "error": str(e)}), 400


@app.route("/api/prodotti/<int:pid>", methods=["PATCH"])
def api_prodotti_update(pid):
    data = request.json
    conn = get_db()
    old = conn.execute("SELECT * FROM prodotti WHERE id=?", (pid,)).fetchone()
    if not old:
        conn.close()
        return jsonify({"ok": False, "error": "Non trovato"}), 404
    old = dict(old)

    # Log storico prezzi se cambiano
    new_acq = data.get("prezzo_acquisto")
    new_ven = data.get("prezzo_vendita")
    if new_acq is not None or new_ven is not None:
        if (new_acq is not None and new_acq != old["prezzo_acquisto"]) or \
           (new_ven is not None and new_ven != old["prezzo_vendita"]):
            now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            conn.execute(
                """INSERT INTO prodotti_prezzi_storico
                   (prodotto_id,prezzo_acquisto_precedente,prezzo_vendita_precedente,
                    prezzo_acquisto_nuovo,prezzo_vendita_nuovo,aggiornato_da,created_at)
                   VALUES (?,?,?,?,?,?,?)""",
                (pid, old["prezzo_acquisto"], old["prezzo_vendita"],
                 new_acq if new_acq is not None else old["prezzo_acquisto"],
                 new_ven if new_ven is not None else old["prezzo_vendita"],
                 data.get("aggiornato_da", ""), now),
            )
            if "data_ultimo_prezzo" not in data:
                data["data_ultimo_prezzo"] = now[:10]

    allowed = ["nome", "categoria", "modello", "trasmissione", "dn",
               "prezzo_acquisto", "prezzo_vendita", "data_ultimo_prezzo",
               "fornitore", "note", "attivo", "ordine"]
    sets, vals = [], []
    for k in allowed:
        if k in data:
            sets.append(k + "=?")
            vals.append(data[k])
    if sets:
        sets.append("updated_at=?")
        vals.append(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        vals.append(pid)
        conn.execute("UPDATE prodotti SET " + ",".join(sets) + " WHERE id=?", vals)
        conn.commit()
    row = conn.execute("SELECT * FROM prodotti WHERE id=?", (pid,)).fetchone()
    conn.close()
    return jsonify({"ok": True, "data": dict(row)})


@app.route("/api/prodotti/<int:pid>", methods=["DELETE"])
def api_prodotti_delete(pid):
    conn = get_db()
    conn.execute("DELETE FROM prodotti WHERE id=?", (pid,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/prodotti/<int:pid>/storico", methods=["GET"])
def api_prodotti_storico(pid):
    conn = get_db()
    rows = conn.execute(
        "SELECT * FROM prodotti_prezzi_storico WHERE prodotto_id=? ORDER BY created_at DESC",
        (pid,),
    ).fetchall()
    conn.close()
    return jsonify({"ok": True, "data": [dict(r) for r in rows]})


# ── Prompt 1: API Modelli apparecchio ────────────────────────────────

@app.route("/api/modelli", methods=["GET"])
def api_modelli_list():
    conn = get_db()
    cat = request.args.get("categoria")
    if cat:
        rows = conn.execute("SELECT * FROM modelli_apparecchio WHERE categoria=? ORDER BY ordine", (cat,)).fetchall()
    else:
        rows = conn.execute("SELECT * FROM modelli_apparecchio ORDER BY ordine").fetchall()
    conn.close()
    return jsonify({"ok": True, "data": [dict(r) for r in rows]})


@app.route("/api/modelli", methods=["POST"])
def api_modelli_create():
    data = request.json
    conn = get_db()
    cur = conn.execute(
        "INSERT INTO modelli_apparecchio (categoria,nome,icona,trasmissione_disponibile,dn_disponibile,tipo_lettura,ordine) VALUES (?,?,?,?,?,?,?)",
        (data.get("categoria"), data.get("nome"), data.get("icona", "thermometer"),
         data.get("trasmissione_disponibile", 0), data.get("dn_disponibile", 0),
         data.get("tipo_lettura"), data.get("ordine", 0)),
    )
    conn.commit()
    row = conn.execute("SELECT * FROM modelli_apparecchio WHERE id=?", (cur.lastrowid,)).fetchone()
    conn.close()
    return jsonify({"ok": True, "data": dict(row)}), 201


@app.route("/api/modelli/<int:mid>", methods=["PATCH"])
def api_modelli_update(mid):
    data = request.json
    conn = get_db()
    allowed = ["nome", "icona", "trasmissione_disponibile", "dn_disponibile", "tipo_lettura", "attivo", "ordine"]
    sets, vals = [], []
    for k in allowed:
        if k in data:
            sets.append(k + "=?")
            vals.append(data[k])
    if sets:
        vals.append(mid)
        conn.execute("UPDATE modelli_apparecchio SET " + ",".join(sets) + " WHERE id=?", vals)
        conn.commit()
    row = conn.execute("SELECT * FROM modelli_apparecchio WHERE id=?", (mid,)).fetchone()
    conn.close()
    return jsonify({"ok": True, "data": dict(row)})


@app.route("/api/modelli/<int:mid>", methods=["DELETE"])
def api_modelli_delete(mid):
    conn = get_db()
    conn.execute("DELETE FROM modelli_apparecchio WHERE id=?", (mid,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


# ── Prompt 1: API Prezzi installazione ───────────────────────────────

@app.route("/api/prezzi-installazione", methods=["GET"])
def api_prezzi_inst_list():
    conn = get_db()
    rows = conn.execute("SELECT * FROM prezzi_installazione ORDER BY id").fetchall()
    conn.close()
    return jsonify({"ok": True, "data": [dict(r) for r in rows]})


@app.route("/api/prezzi-installazione/<int:pid>", methods=["PATCH"])
def api_prezzi_inst_update(pid):
    data = request.json
    conn = get_db()
    allowed = ["descrizione", "prezzo_base", "unita", "note"]
    sets, vals = [], []
    for k in allowed:
        if k in data:
            sets.append(k + "=?")
            vals.append(data[k])
    if sets:
        sets.append("updated_at=?")
        vals.append(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        vals.append(pid)
        conn.execute("UPDATE prezzi_installazione SET " + ",".join(sets) + " WHERE id=?", vals)
        conn.commit()
    row = conn.execute("SELECT * FROM prezzi_installazione WHERE id=?", (pid,)).fetchone()
    conn.close()
    return jsonify({"ok": True, "data": dict(row)})


# ── Prompt 2: Generatore Offerte IA ──────────────────────────────────

@app.route("/generatore")
def generatore_page():
    return render_template("generatore.html")


@app.route("/api/templates", methods=["GET"])
def api_templates_list():
    templates = []
    for f in UPLOADS_DIR.glob("*.docx"):
        templates.append({
            "filename": f.name,
            "size": f.stat().st_size,
            "modified": datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d %H:%M"),
            "tipo": "E40" if "E40" in f.name else ("Q55" if "Q5.5" in f.name or "Q55" in f.name else "altro"),
        })
    return jsonify({"ok": True, "data": templates})


@app.route("/api/generatore/crea", methods=["POST"])
def api_generatore_crea():
    data = request.json
    conn = get_db()
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # 1. Validate
    required = ["nome_studio", "cond_via", "cond_comune", "natura", "tipo_offerta", "agente_id"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        conn.close()
        return jsonify({"ok": False, "error": "Campi mancanti: " + ", ".join(missing)}), 400

    # 2. Numero progressivo
    cfg = read_config()
    numero = cfg["prossimo_numero"]
    cfg["prossimo_numero"] = numero + 1
    write_config(cfg)

    # 3. Cliente
    nome_studio = data["nome_studio"].strip()
    cliente_id = None
    if data.get("cliente_id"):
        cliente_id = int(data["cliente_id"])
    else:
        existing = conn.execute("SELECT id FROM clienti WHERE LOWER(nome_studio)=LOWER(?)", (nome_studio,)).fetchone()
        if existing:
            cliente_id = existing[0]
        elif data.get("salva_anagrafica", True):
            cur = conn.execute(
                "INSERT INTO clienti (nome_studio,via,cap,citta,email,telefono,referente,tipo_cliente,data_inserimento) VALUES (?,?,?,?,?,?,?,?,?)",
                (nome_studio, data.get("cliente_via", ""), data.get("cliente_cap", ""),
                 data.get("cliente_citta", ""), data.get("cliente_email", ""),
                 data.get("cliente_telefono", ""), data.get("cliente_referente", ""),
                 data.get("tipo_cliente", "Amministratore"), now),
            )
            cliente_id = cur.lastrowid

    # 4. Oggetto/Condominio
    oggetto_id = None
    if data.get("oggetto_id"):
        oggetto_id = int(data["oggetto_id"])
    else:
        cur = conn.execute(
            "INSERT INTO oggetti (cliente_id,nome,via,civico,comune,provincia,cap,agente_id,natura,n_unita,n_scale,created_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
            (cliente_id, data.get("cond_nome", ""), data["cond_via"],
             data.get("cond_civico", ""), data["cond_comune"],
             data.get("cond_provincia", ""), data.get("cond_cap", ""),
             int(data["agente_id"]), data["natura"],
             data.get("n_unita"), data.get("n_scale"), now),
        )
        oggetto_id = cur.lastrowid

    # 5. Calcola importi
    apparecchi = data.get("apparecchi", [])
    importo_fornitura = 0
    for ap in apparecchi:
        importo_fornitura += (ap.get("prezzo_vendita", 0) or 0) * (ap.get("quantita", 0) or 0)

    centralizzazione = data.get("centralizzazione", {})
    if centralizzazione.get("attiva") and centralizzazione.get("tipo_fornitura") == "vendita":
        importo_fornitura += (centralizzazione.get("prezzo_unitario", 0) or 0) * (centralizzazione.get("quantita", 1) or 1)

    servizi = data.get("servizi", {})
    importo_annuo = 0
    lettura = servizi.get("lettura", {})
    if lettura.get("prezzo", 0):
        importo_annuo += (lettura.get("prezzo", 0) or 0) * (lettura.get("quantita", 0) or 0)
    care = servizi.get("care", {})
    if care.get("attivo") and care.get("prezzo", 0):
        importo_annuo += (care.get("prezzo", 0) or 0) * (care.get("quantita", 0) or 0)

    # Determine template
    template_file = data.get("template", "")
    template_key = "E40" if "E40" in template_file else ("Q55" if "Q5.5" in template_file or "Q55" in template_file else "E40")

    # 6. INSERT offerta
    cur = conn.execute(
        """INSERT INTO offerte (numero,nome_studio,nome_condominio,via,cap,citta,
           riferimento,template,prezzo_fornitura,prezzo_care,canone_lettura,
           modalita,importo,importo_servizio_annuo,stato,email_studio,
           data_creazione,agente_id,oggetto_id,tipo_offerta,natura,
           is_accordo_quadro,stato_versione,versione)
           VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (numero, nome_studio, data.get("cond_nome", ""),
         data["cond_via"], data.get("cond_cap", ""), data["cond_comune"],
         data.get("riferimento", ""), template_key,
         lettura.get("prezzo") if lettura else None,
         care.get("prezzo") if care.get("attivo") else None,
         lettura.get("prezzo") if lettura else None,
         data.get("modalita", "vendita"),
         importo_fornitura if importo_fornitura else None,
         importo_annuo if importo_annuo else None,
         "richiamato", data.get("cliente_email", ""), now,
         int(data["agente_id"]), oggetto_id,
         data["tipo_offerta"], data["natura"],
         1 if data.get("is_accordo_quadro") else 0, "attiva", "A"),
    )
    offerta_id = cur.lastrowid

    # 7. INSERT righe
    ordine = 0
    for ap in apparecchi:
        ordine += 1
        conn.execute(
            "INSERT INTO offerte_righe (offerta_id,descrizione,tipo_riga,prezzo_unitario,quantita,quantita_stimata,totale_riga,ordine) VALUES (?,?,?,?,?,?,?,?)",
            (offerta_id, ap.get("modello", "") + " " + ap.get("categoria", ""),
             "fornitura", ap.get("prezzo_vendita", 0), ap.get("quantita", 0),
             1 if ap.get("stimata") else 0,
             (ap.get("prezzo_vendita", 0) or 0) * (ap.get("quantita", 0) or 0), ordine),
        )
    if centralizzazione.get("attiva") and centralizzazione.get("tipo_fornitura") == "vendita":
        ordine += 1
        conn.execute(
            "INSERT INTO offerte_righe (offerta_id,descrizione,tipo_riga,prezzo_unitario,quantita,totale_riga,ordine) VALUES (?,?,?,?,?,?,?)",
            (offerta_id, "Concentratore " + centralizzazione.get("modello", ""),
             "fornitura", centralizzazione.get("prezzo_unitario", 0),
             centralizzazione.get("quantita", 1),
             (centralizzazione.get("prezzo_unitario", 0) or 0) * (centralizzazione.get("quantita", 1) or 1), ordine),
        )
    if lettura.get("prezzo", 0):
        ordine += 1
        conn.execute(
            "INSERT INTO offerte_righe (offerta_id,descrizione,tipo_riga,prezzo_unitario,quantita,totale_riga,ordine) VALUES (?,?,?,?,?,?,?)",
            (offerta_id, "Lettura consumi " + lettura.get("tipo", "RK"),
             "servizio_annuo", lettura.get("prezzo", 0), lettura.get("quantita", 0),
             (lettura.get("prezzo", 0) or 0) * (lettura.get("quantita", 0) or 0), ordine),
        )
    if care.get("attivo") and care.get("prezzo", 0):
        ordine += 1
        conn.execute(
            "INSERT INTO offerte_righe (offerta_id,descrizione,tipo_riga,prezzo_unitario,quantita,totale_riga,ordine) VALUES (?,?,?,?,?,?,?)",
            (offerta_id, "Ulteria Care", "care", care.get("prezzo", 0), care.get("quantita", 0),
             (care.get("prezzo", 0) or 0) * (care.get("quantita", 0) or 0), ordine),
        )

    conn.commit()

    # 8-11. DOCX Generation
    slug = studio_slug(nome_studio)
    tipo = template_key
    anno = datetime.now().strftime("%Y")
    folder_name = f"{numero}_{slug}_{tipo}"
    dest_dir = OUTPUT_DIR / anno / folder_name
    dest_dir.mkdir(parents=True, exist_ok=True)
    file_base = f"ULTERIA_{numero}_{slug}_{tipo}"
    docx_name = f"{file_base}.docx"
    docx_path = dest_dir / docx_name

    # Find template source
    src = None
    if template_file and (UPLOADS_DIR / template_file).exists():
        src = UPLOADS_DIR / template_file
    elif template_key in TEMPLATE_MAP:
        src = UPLOADS_DIR / TEMPLATE_MAP[template_key]

    docx_rel = None
    pdf_rel = None
    pdf_error = False

    if src and src.exists():
        shutil.copy2(str(src), str(docx_path))
        doc = Document(str(docx_path))

        via_full = data["cond_via"] + (" " + data.get("cond_civico", "") if data.get("cond_civico") else "")
        dat_str = datetime.now().strftime("%d/%m/%Y")
        agente_row = conn.execute("SELECT nome,cognome FROM agenti WHERE id=?", (int(data["agente_id"]),)).fetchone()
        agente_nome = (agente_row["nome"] + " " + agente_row["cognome"]) if agente_row else ""

        # Ripartitore info
        rip = next((a for a in apparecchi if a.get("categoria") == "ripartitore"), None)
        ca = next((a for a in apparecchi if a.get("categoria") == "contatore_acqua"), None)
        cc = next((a for a in apparecchi if a.get("categoria") == "contatore_calore"), None)
        tipo_lettura = "RD" if cc else "RK"

        # Replace placeholders — existing style «»
        replace_runs(doc.paragraphs, "\u00abSTU\u00bb", nome_studio)
        replace_runs(doc.paragraphs, "\u00abVIA\u00bb", via_full)
        replace_runs(doc.paragraphs, "\u00abDAT\u00bb", dat_str)

        fix_stu_via_paragraph(doc, nome_studio, via_full)

        # All sections XML
        for section in doc.sections:
            for part in [section.header, section.footer, section.first_page_header,
                         section.first_page_footer, section.even_page_header, section.even_page_footer]:
                if part:
                    try:
                        replace_xml(part._element, "\u00abNR\u00bb", str(numero))
                        replace_xml(part._element, "\u00abSTU\u00bb", nome_studio)
                    except Exception:
                        pass

        # Body XML — all placeholders
        body_el = doc.element.body
        replace_xml(body_el, "\u00abNR\u00bb", str(numero))
        replace_xml(body_el, "\u00abSTU\u00bb", nome_studio)
        replace_xml(body_el, "\u00abVIA\u00bb", via_full)
        replace_xml(body_el, "\u00abDAT\u00bb", dat_str)
        replace_xml(body_el, "\u00abPFO\u00bb", format_eur(importo_fornitura) if importo_fornitura else "\u2014")
        replace_xml(body_el, "\u00abPCA\u00bb", format_eur(care.get("prezzo", 0)) if care.get("attivo") else "\u2014")
        replace_xml(body_el, "\u00abPCL\u00bb", format_eur(lettura.get("prezzo", 0)) if lettura.get("prezzo") else "\u2014")
        replace_xml(body_el, "\u00abMOD\u00bb", rip["modello"] if rip else (ca["modello"] if ca else ""))

        doc.save(str(docx_path))
        docx_rel = f"/output/{anno}/{folder_name}/{docx_name}"

        # PDF
        try:
            import pythoncom
            pythoncom.CoInitialize()
            try:
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                word.DisplayAlerts = False
                doc_com = word.Documents.Open(str(docx_path.resolve()))
                pdf_path = dest_dir / f"{file_base}.pdf"
                doc_com.SaveAs2(str(pdf_path.resolve()), FileFormat=17)
                doc_com.Close(False)
                word.Quit()
                pdf_rel = f"/output/{anno}/{folder_name}/{file_base}.pdf"
            finally:
                pythoncom.CoUninitialize()
        except Exception as e:
            pdf_error = True
            app.logger.warning(f"PDF conversion failed: {e}")

    # 12. Update offerta paths
    conn.execute("UPDATE offerte SET path_docx=?, path_pdf=? WHERE id=?", (docx_rel, pdf_rel, offerta_id))

    # Timeline event
    conn.execute(
        "INSERT INTO timeline_eventi (tipo_evento,descrizione,offerta_id,oggetto_id,created_at) VALUES (?,?,?,?,?)",
        ("offerta_creata", f"Offerta {numero} generata via Generatore IA", offerta_id, oggetto_id, now),
    )
    conn.commit()
    conn.close()

    return jsonify({
        "ok": True,
        "offerta_id": offerta_id,
        "numero_offerta": str(numero),
        "path_docx": docx_rel,
        "path_pdf": pdf_rel,
        "pdf_error": pdf_error,
    })


# ── Prompt 3: Fogli Costi API ────────────────────────────────────────

@app.route("/api/fogli-costi/by-oggetto/<int:oid>", methods=["GET"])
def api_fogli_costi_by_oggetto(oid):
    conn = get_db()
    fc = conn.execute("SELECT * FROM fogli_costi WHERE oggetto_id=? ORDER BY id DESC LIMIT 1", (oid,)).fetchone()
    extras = []
    if fc:
        extras = conn.execute("SELECT * FROM foglio_costi_extra WHERE foglio_costi_id=? ORDER BY id", (fc["id"],)).fetchall()
    # Get offerta righe for this oggetto
    righe = conn.execute(
        "SELECT or2.*, o.numero as offerta_numero FROM offerte_righe or2 "
        "JOIN offerte o ON or2.offerta_id = o.id "
        "WHERE o.oggetto_id=? AND o.stato_versione='attiva' ORDER BY or2.ordine", (oid,)
    ).fetchall()
    # Get prodotti for price lookup
    prodotti = conn.execute("SELECT codice,nome,modello,prezzo_acquisto FROM prodotti WHERE attivo=1").fetchall()
    # Get prezzi installazione
    prezzi_inst = conn.execute("SELECT * FROM prezzi_installazione").fetchall()
    conn.close()
    return jsonify({
        "ok": True,
        "data": {
            "foglio": dict(fc) if fc else None,
            "extras": [dict(e) for e in extras],
            "righe_offerta": [dict(r) for r in righe],
            "prodotti": [dict(p) for p in prodotti],
            "prezzi_installazione": [dict(p) for p in prezzi_inst],
        }
    })


@app.route("/api/fogli-costi", methods=["POST"])
def api_fogli_costi_create():
    data = request.json
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn = get_db()
    cur = conn.execute(
        """INSERT INTO fogli_costi (oggetto_id, offerta_id, costo_apparecchi,
           costo_installazione_idraulica, costo_installazione_elettrica,
           costo_concentratori, costo_materiali_extra, note_costi, totale_costi,
           ricavo_fornitura, ricavo_servizio_annuo, k_moltiplicatore,
           margine_euro, margine_percentuale,
           provvigione1_nome, provvigione1_percentuale, provvigione1_euro,
           provvigione2_nome, provvigione2_percentuale, provvigione2_euro,
           provvigione3_nome, provvigione3_percentuale, provvigione3_euro,
           netto_finale, created_at, updated_at)
           VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (data.get("oggetto_id"), data.get("offerta_id"),
         data.get("costo_apparecchi", 0), data.get("costo_installazione_idraulica", 0),
         data.get("costo_installazione_elettrica", 0), data.get("costo_concentratori", 0),
         data.get("costo_materiali_extra", 0), data.get("note_costi", ""),
         data.get("totale_costi", 0), data.get("ricavo_fornitura", 0),
         data.get("ricavo_servizio_annuo", 0), data.get("k_moltiplicatore", 1.0),
         data.get("margine_euro", 0), data.get("margine_percentuale", 0),
         data.get("provvigione1_nome"), data.get("provvigione1_percentuale", 0), data.get("provvigione1_euro", 0),
         data.get("provvigione2_nome"), data.get("provvigione2_percentuale", 0), data.get("provvigione2_euro", 0),
         data.get("provvigione3_nome"), data.get("provvigione3_percentuale", 0), data.get("provvigione3_euro", 0),
         data.get("netto_finale", 0), now, now),
    )
    conn.commit()
    row = conn.execute("SELECT * FROM fogli_costi WHERE id=?", (cur.lastrowid,)).fetchone()
    conn.close()
    return jsonify({"ok": True, "data": dict(row)}), 201


@app.route("/api/fogli-costi/<int:fid>", methods=["PATCH"])
def api_fogli_costi_update(fid):
    data = request.json
    conn = get_db()
    allowed = [
        "scenario", "installatore_idraulico",
        "costo_apparecchi", "costo_installazione_idraulica", "costo_installazione_elettrica",
        "costo_concentratori", "costo_materiali_extra", "note_costi", "totale_costi",
        "ricavo_fornitura", "ricavo_servizio_annuo", "k_moltiplicatore",
        "margine_euro", "margine_percentuale",
        "provvigione1_nome", "provvigione1_percentuale", "provvigione1_euro",
        "provvigione2_nome", "provvigione2_percentuale", "provvigione2_euro",
        "provvigione3_nome", "provvigione3_percentuale", "provvigione3_euro",
        "netto_finale",
        "cont_riscaldamento", "cont_riscaldamento_trasmissione", "cont_riscaldamento_dn", "cont_riscaldamento_costo",
        "cont_hc", "cont_hc_trasmissione", "cont_hc_dn", "cont_hc_costo",
        "cont_raffrescamento", "cont_raffrescamento_trasmissione", "cont_raffrescamento_dn", "cont_raffrescamento_costo",
        "cont_acqua_calda", "cont_acqua_calda_trasmissione", "cont_acqua_calda_dn", "cont_acqua_calda_costo",
        "cont_acqua_fredda", "cont_acqua_fredda_trasmissione", "cont_acqua_fredda_dn", "cont_acqua_fredda_costo",
        "cont_acqua_ricircolo", "cont_acqua_ricircolo_costo",
        "cont_acqua_duale", "cont_acqua_duale_costo",
        "inst_cont_calore", "inst_cont_acqua_calda", "inst_cont_acqua_fredda", "inst_modifiche_idrauliche",
        "costo_valvola_zona", "costo_attuatore", "costo_produzione_modulo", "costo_opere_idrauliche_extra",
        "costo_trasformatore", "costo_rele", "costo_elettricista", "costo_collegamenti_elettrici",
        "costo_valvola_intercettazione",
        "centr_famiglia", "centr_modello", "centr_pezzi", "centr_costo_acquisto",
        "centr_costo_installazione", "centr_pw_router", "centr_collaudo",
        "n_radiatori", "costo_kit_valvola", "costo_montaggio_valvola",
        "costo_apparecchio_ripartitore", "costo_extra_trasporto",
        "servizio_lettura_tipo", "servizio_lettura_cad", "servizio_lettura_totale",
    ]
    sets, vals = [], []
    for k in allowed:
        if k in data:
            sets.append(k + "=?")
            vals.append(data[k])
    if sets:
        sets.append("updated_at=?")
        vals.append(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        vals.append(fid)
        conn.execute("UPDATE fogli_costi SET " + ",".join(sets) + " WHERE id=?", vals)
        conn.commit()
    row = conn.execute("SELECT * FROM fogli_costi WHERE id=?", (fid,)).fetchone()
    conn.close()
    return jsonify({"ok": True, "data": dict(row)})


@app.route("/api/fogli-costi/<int:fid>/extra", methods=["POST"])
def api_fogli_costi_add_extra(fid):
    data = request.json
    conn = get_db()
    cur = conn.execute(
        "INSERT INTO foglio_costi_extra (foglio_costi_id, descrizione, quantita, prezzo_unitario, totale) VALUES (?,?,?,?,?)",
        (fid, data.get("descrizione", ""), data.get("quantita", 1),
         data.get("prezzo_unitario", 0), data.get("totale", 0)),
    )
    conn.commit()
    row = conn.execute("SELECT * FROM foglio_costi_extra WHERE id=?", (cur.lastrowid,)).fetchone()
    conn.close()
    return jsonify({"ok": True, "data": dict(row)}), 201


@app.route("/api/fogli-costi/extra/<int:eid>", methods=["DELETE"])
def api_fogli_costi_del_extra(eid):
    conn = get_db()
    conn.execute("DELETE FROM foglio_costi_extra WHERE id=?", (eid,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/offerte/<int:oid>/righe", methods=["GET"])
def api_offerte_righe(oid):
    conn = get_db()
    rows = conn.execute("SELECT * FROM offerte_righe WHERE offerta_id=? ORDER BY ordine", (oid,)).fetchall()
    conn.close()
    return jsonify({"ok": True, "data": [dict(r) for r in rows]})


@app.route("/api/attivita/badge_count", methods=["GET"])
def api_attivita_badge_count():
    conn = get_db()
    count = conn.execute(
        "SELECT COUNT(*) FROM attivita WHERE stato='aperta' AND date(data_scadenza) <= date('now')"
    ).fetchone()[0]
    conn.close()
    return jsonify({"ok": True, "count": count})


# ── Prompt 4: Segnalatori ────────────────────────────────────────────

@app.route("/segnalatori")
def segnalatori_page():
    return render_template("segnalatori.html")


@app.route("/api/segnalatori", methods=["GET"])
def api_segnalatori_list():
    conn = get_db()
    q = request.args.get("q", "").strip()
    sql = "SELECT s.*, COUNT(os.id) as n_offerte, COALESCE(SUM(os.provvigione_euro),0) as tot_provvigioni, COALESCE(SUM(CASE WHEN os.stato_pagamento='da_pagare' THEN os.provvigione_euro ELSE 0 END),0) as da_pagare FROM segnalatori s LEFT JOIN offerte_segnalatori os ON s.id=os.segnalatore_id"
    params = []
    if q:
        sql += " WHERE s.nome LIKE ?"
        params.append("%" + q + "%")
    sql += " GROUP BY s.id ORDER BY s.nome"
    rows = conn.execute(sql, params).fetchall()
    conn.close()
    return jsonify({"ok": True, "data": [dict(r) for r in rows]})


@app.route("/api/segnalatori", methods=["POST"])
def api_segnalatori_create():
    data = request.json
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn = get_db()
    cur = conn.execute(
        "INSERT INTO segnalatori (nome,tipo,azienda,email,telefono,provvigione_default_pct,note,created_at) VALUES (?,?,?,?,?,?,?,?)",
        (data.get("nome", ""), data.get("tipo", "segnalatore"), data.get("azienda", ""),
         data.get("email", ""), data.get("telefono", ""), data.get("provvigione_default_pct", 0),
         data.get("note", ""), now),
    )
    conn.commit()
    row = conn.execute("SELECT * FROM segnalatori WHERE id=?", (cur.lastrowid,)).fetchone()
    conn.close()
    return jsonify({"ok": True, "data": dict(row)}), 201


@app.route("/api/segnalatori/<int:sid>", methods=["PATCH"])
def api_segnalatori_update(sid):
    data = request.json
    conn = get_db()
    allowed = ["nome", "tipo", "azienda", "email", "telefono", "provvigione_default_pct", "note", "attivo"]
    sets, vals = [], []
    for k in allowed:
        if k in data:
            sets.append(k + "=?")
            vals.append(data[k])
    if sets:
        vals.append(sid)
        conn.execute("UPDATE segnalatori SET " + ",".join(sets) + " WHERE id=?", vals)
        conn.commit()
    row = conn.execute("SELECT * FROM segnalatori WHERE id=?", (sid,)).fetchone()
    conn.close()
    return jsonify({"ok": True, "data": dict(row)})


@app.route("/api/segnalatori/<int:sid>", methods=["GET"])
def api_segnalatori_detail(sid):
    conn = get_db()
    seg = conn.execute("SELECT * FROM segnalatori WHERE id=?", (sid,)).fetchone()
    if not seg:
        conn.close()
        return jsonify({"ok": False, "error": "Non trovato"}), 404
    offerte = conn.execute(
        "SELECT os.*, o.numero, o.nome_studio, o.nome_condominio, o.importo, o.stato "
        "FROM offerte_segnalatori os JOIN offerte o ON os.offerta_id=o.id "
        "WHERE os.segnalatore_id=? ORDER BY os.created_at DESC", (sid,)
    ).fetchall()
    conn.close()
    return jsonify({"ok": True, "data": {"segnalatore": dict(seg), "offerte": [dict(o) for o in offerte]}})


@app.route("/api/segnalatori/search", methods=["GET"])
def api_segnalatori_search():
    q = request.args.get("q", "").strip()
    conn = get_db()
    rows = conn.execute(
        "SELECT id, nome, tipo, provvigione_default_pct FROM segnalatori WHERE attivo=1 AND nome LIKE ? LIMIT 10",
        ("%" + q + "%",),
    ).fetchall()
    conn.close()
    return jsonify({"ok": True, "data": [dict(r) for r in rows]})


@app.route("/api/offerte/<int:oid>/segnalatore", methods=["PATCH"])
def api_offerte_segnalatore(oid):
    data = request.json
    conn = get_db()
    seg_id = data.get("segnalatore_id")
    pct = data.get("provvigione_pct", 0)
    conn.execute("UPDATE offerte SET segnalatore_id=? WHERE id=?", (seg_id, oid))
    off = conn.execute("SELECT importo FROM offerte WHERE id=?", (oid,)).fetchone()
    importo_base = (off["importo"] or 0) if off else 0
    provv_euro = importo_base * pct / 100
    existing = conn.execute("SELECT id FROM offerte_segnalatori WHERE offerta_id=? AND segnalatore_id=?", (oid, seg_id)).fetchone()
    if existing:
        conn.execute("UPDATE offerte_segnalatori SET provvigione_pct=?, importo_base=?, provvigione_euro=? WHERE id=?",
                     (pct, importo_base, provv_euro, existing["id"]))
    else:
        conn.execute("INSERT INTO offerte_segnalatori (offerta_id,segnalatore_id,provvigione_pct,importo_base,provvigione_euro) VALUES (?,?,?,?,?)",
                     (oid, seg_id, pct, importo_base, provv_euro))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/api/offerte-segnalatori/<int:osid>/paga", methods=["PATCH"])
def api_offerte_segnalatori_paga(osid):
    data = request.json
    now = datetime.now().strftime("%Y-%m-%d")
    conn = get_db()
    conn.execute("UPDATE offerte_segnalatori SET stato_pagamento='pagato', data_pagamento=? WHERE id=?", (data.get("data_pagamento", now), osid))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


# ── Prompt 4: Cambio Agente ──────────────────────────────────────────

@app.route("/api/oggetti/<int:oid>/cambio-agente", methods=["POST"])
def api_oggetti_cambio_agente(oid):
    data = request.json
    conn = get_db()
    obj = conn.execute("SELECT agente_id FROM oggetti WHERE id=?", (oid,)).fetchone()
    old_id = obj["agente_id"] if obj else None
    new_id = data.get("agente_id")
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn.execute(
        "INSERT INTO oggetti_agenti_storico (oggetto_id,agente_id_precedente,agente_id_nuovo,motivo,data_cambio,effettuato_da) VALUES (?,?,?,?,?,?)",
        (oid, old_id, new_id, data.get("motivo", ""), data.get("data_cambio", now[:10]), data.get("effettuato_da", "")),
    )
    conn.execute("UPDATE oggetti SET agente_id=?, updated_at=? WHERE id=?", (new_id, now, oid))
    conn.execute(
        "INSERT INTO timeline_eventi (tipo_evento,descrizione,oggetto_id,utente,created_at) VALUES (?,?,?,?,?)",
        ("cambio_agente", "Agente cambiato", oid, data.get("effettuato_da", ""), now),
    )
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


# ── Prompt 4: Dashboard Servizi ──────────────────────────────────────

@app.route("/api/dashboard/servizi", methods=["GET"])
def api_dashboard_servizi():
    conn = get_db()
    agente_id = request.args.get("agente_id")
    sql = "SELECT tipo_servizio, COUNT(*) as count, SUM(CASE WHEN stato='preso_lavoro' THEN 1 ELSE 0 END) as prese, SUM(CASE WHEN stato='preso_lavoro' THEN importo_servizio_annuo ELSE 0 END) as canone_annuo FROM offerte WHERE tipo_servizio IS NOT NULL AND stato_versione='attiva'"
    params = []
    if agente_id:
        sql += " AND agente_id=?"
        params.append(int(agente_id))
    sql += " GROUP BY tipo_servizio"
    rows = conn.execute(sql, params).fetchall()
    conn.close()
    return jsonify({"ok": True, "data": [dict(r) for r in rows]})


# ── Prompt 4: Export CSV ─────────────────────────────────────────────

@app.route("/api/export/offerte", methods=["GET"])
def api_export_offerte():
    from flask import Response
    conn = get_db()
    sql = "SELECT o.*, a.nome as ag_nome, a.cognome as ag_cognome, s.nome as seg_nome FROM offerte o LEFT JOIN agenti a ON o.agente_id=a.id LEFT JOIN segnalatori s ON o.segnalatore_id=s.id WHERE o.stato_versione='attiva'"
    params = []
    if request.args.get("agente_id"):
        sql += " AND o.agente_id=?"
        params.append(int(request.args["agente_id"]))
    if request.args.get("dal"):
        sql += " AND o.data_creazione>=?"
        params.append(request.args["dal"])
    if request.args.get("al"):
        sql += " AND o.data_creazione<=?"
        params.append(request.args["al"] + " 23:59:59")
    sql += " ORDER BY o.numero DESC"
    rows = conn.execute(sql, params).fetchall()
    conn.close()

    header = "NUMERO;DATA;AGENTE;CLIENTE;VIA_CONDOMINIO;COMUNE;PROVINCIA;TIPO_OFFERTA;NATURA;TIPO_SERVIZIO;VALORE_FORNITURA;VALORE_SERVIZIO_ANNUO;STATO;SEGNALATORE;PROVVIGIONE_SEGNALATORE_PCT;IS_ACCORDO_QUADRO;NOTE\n"
    lines = []
    for r in rows:
        ag = (r["ag_nome"] or "") + " " + (r["ag_cognome"] or "")
        d = r["data_creazione"][:10] if r["data_creazione"] else ""
        if d and "-" in d:
            parts = d.split("-")
            d = parts[2] + "/" + parts[1] + "/" + parts[0]
        line = ";".join([
            str(r["numero"] or ""), d, ag.strip(), r["nome_studio"] or "",
            r["via"] or "", r["citta"] or "", "", r["tipo_offerta"] or "",
            r["natura"] or "", r["tipo_servizio"] or "",
            str(r["importo"] or ""), str(r["importo_servizio_annuo"] or ""),
            r["stato"] or "", r["seg_nome"] or "", "",
            str(r["is_accordo_quadro"] or 0), (r["note"] or "").replace(";", ",")
        ])
        lines.append(line)

    csv_content = "\ufeff" + header + "\n".join(lines)
    return Response(csv_content, mimetype="text/csv",
                    headers={"Content-Disposition": "attachment;filename=offerte_ulteria.csv"})


@app.route("/api/export/clienti", methods=["GET"])
def api_export_clienti():
    from flask import Response
    conn = get_db()
    rows = conn.execute("SELECT * FROM clienti ORDER BY nome_studio").fetchall()
    conn.close()
    header = "NOME_STUDIO;REFERENTE;VIA;CIVICO;COMUNE;CAP;PROVINCIA;EMAIL;TELEFONO;TIPO_CLIENTE;SETTORE;NOTE\n"
    lines = []
    for r in rows:
        line = ";".join([
            r["nome_studio"] or "", r["referente"] or "", r["via"] or "", "",
            r["citta"] or "", r["cap"] or "", "", r["email"] or "",
            r["telefono"] or "", r["tipo_cliente"] or "", r["settore"] or "",
            (r["note"] or "").replace(";", ",")
        ])
        lines.append(line)
    csv_content = "\ufeff" + header + "\n".join(lines)
    return Response(csv_content, mimetype="text/csv",
                    headers={"Content-Disposition": "attachment;filename=clienti_ulteria.csv"})


@app.route("/api/import/template/offerte", methods=["GET"])
def api_import_template_offerte():
    from flask import Response
    csv = "\ufeffNUMERO;DATA;AGENTE;CLIENTE;VIA_CONDOMINIO;COMUNE;PROVINCIA;TIPO_OFFERTA;NATURA;TIPO_SERVIZIO;VALORE_FORNITURA;VALORE_SERVIZIO_ANNUO;STATO;SEGNALATORE;PROVVIGIONE_SEGNALATORE_PCT;IS_ACCORDO_QUADRO;NOTE\n"
    csv += "#26001;01/01/2026;FB;Studio Rossi;Via Roma 1;Monza;;installazione;nuovo;RK;12400;1800;richiamato;Piaggi;3;0;Esempio\n"
    return Response(csv, mimetype="text/csv", headers={"Content-Disposition": "attachment;filename=template_offerte.csv"})


@app.route("/api/import/template/clienti", methods=["GET"])
def api_import_template_clienti():
    from flask import Response
    csv = "\ufeffNOME_STUDIO;REFERENTE;VIA;CIVICO;COMUNE;CAP;PROVINCIA;EMAIL;TELEFONO;TIPO_CLIENTE;SETTORE;NOTE\n"
    csv += "#Studio Esempio;Mario Rossi;Via Roma;1;Monza;20900;;info@studio.it;039123456;Amministratore;amministratori;Esempio\n"
    return Response(csv, mimetype="text/csv", headers={"Content-Disposition": "attachment;filename=template_clienti.csv"})


@app.route("/api/import/offerte", methods=["POST"])
def api_import_offerte():
    import csv as csvmod
    import io
    f = request.files.get("file")
    if not f:
        return jsonify({"ok": False, "error": "Nessun file"}), 400
    content = f.read().decode("utf-8-sig")
    reader = csvmod.DictReader(io.StringIO(content), delimiter=";")

    conn = get_db()
    agenti_map = {}
    for a in conn.execute("SELECT id, nome, cognome FROM agenti").fetchall():
        ini = (a["nome"][0] if a["nome"] else "") + (a["cognome"][0] if a["cognome"] else "")
        agenti_map[ini.upper()] = a["id"]
        agenti_map[(a["nome"] + " " + a["cognome"]).strip().upper()] = a["id"]

    imported = 0
    skipped = 0
    errors = []
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    for i, row in enumerate(reader):
        try:
            if not row.get("CLIENTE") and not row.get("NUMERO"):
                skipped += 1
                continue
            # Parse agent (handle "FB / PIAGGI")
            ag_str = (row.get("AGENTE") or "").strip()
            agente_id = None
            segnalatore_nome = None
            if "/" in ag_str:
                parts = ag_str.split("/")
                ag_str = parts[0].strip()
                segnalatore_nome = parts[1].strip() if len(parts) > 1 else None
            agente_id = agenti_map.get(ag_str.upper())

            # Parse importo (handle Excel formulas)
            importo = row.get("VALORE_FORNITURA", "") or ""
            if str(importo).startswith("="):
                importo = 0
            else:
                try:
                    importo = float(str(importo).replace(",", ".")) if importo else None
                except ValueError:
                    importo = None

            importo_annuo = row.get("VALORE_SERVIZIO_ANNUO", "") or ""
            try:
                importo_annuo = float(str(importo_annuo).replace(",", ".")) if importo_annuo else None
            except ValueError:
                importo_annuo = None

            # Parse date
            data = row.get("DATA", "")
            if data and "/" in data:
                parts = data.split("/")
                if len(parts) == 3:
                    data = parts[2] + "-" + parts[1] + "-" + parts[0]

            # Find or create client
            nome_studio = (row.get("CLIENTE") or "").strip()
            cliente = conn.execute("SELECT id FROM clienti WHERE LOWER(nome_studio)=LOWER(?)", (nome_studio,)).fetchone()
            if not cliente and nome_studio:
                conn.execute("INSERT INTO clienti (nome_studio,citta,data_inserimento) VALUES (?,?,?)",
                             (nome_studio, row.get("COMUNE", ""), now))

            numero = row.get("NUMERO", "")
            try:
                numero = int(str(numero).replace("X", "").replace("x", "").strip()) if numero else None
            except ValueError:
                numero = None

            conn.execute(
                """INSERT INTO offerte (numero,nome_studio,via,citta,tipo_offerta,natura,tipo_servizio,
                   importo,importo_servizio_annuo,stato,agente_id,data_creazione,
                   is_accordo_quadro,stato_versione,versione,note)
                   VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (numero, nome_studio, row.get("VIA_CONDOMINIO", ""), row.get("COMUNE", ""),
                 row.get("TIPO_OFFERTA", "installazione"), row.get("NATURA", "nuovo"),
                 row.get("TIPO_SERVIZIO") or None, importo, importo_annuo,
                 row.get("STATO", "richiamato"), agente_id,
                 (data + " 00:00:00") if data else now,
                 int(row.get("IS_ACCORDO_QUADRO", 0) or 0), "attiva", "A",
                 row.get("NOTE", "")),
            )

            # Handle segnalatore
            if segnalatore_nome:
                seg = conn.execute("SELECT id FROM segnalatori WHERE nome LIKE ?", ("%" + segnalatore_nome + "%",)).fetchone()
                if seg:
                    conn.execute("UPDATE offerte SET segnalatore_id=? WHERE rowid=last_insert_rowid()", (seg["id"],))

            imported += 1
        except Exception as e:
            skipped += 1
            errors.append({"riga": i + 2, "errore": str(e)})

    conn.commit()
    conn.close()
    return jsonify({"ok": True, "data": {"imported": imported, "skipped": skipped, "errors": errors}})


@app.route("/api/import/clienti", methods=["POST"])
def api_import_clienti():
    import csv as csvmod
    import io
    f = request.files.get("file")
    if not f:
        return jsonify({"ok": False, "error": "Nessun file"}), 400
    content = f.read().decode("utf-8-sig")
    reader = csvmod.DictReader(io.StringIO(content), delimiter=";")

    conn = get_db()
    imported = 0
    updated = 0
    skipped = 0
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    for i, row in enumerate(reader):
        try:
            nome = (row.get("NOME_STUDIO") or "").strip()
            if not nome:
                skipped += 1
                continue
            existing = conn.execute("SELECT id FROM clienti WHERE LOWER(nome_studio)=LOWER(?)", (nome,)).fetchone()
            if existing:
                conn.execute(
                    "UPDATE clienti SET via=COALESCE(NULLIF(?,'')),citta=COALESCE(NULLIF(?,'')),cap=COALESCE(NULLIF(?,'')),email=COALESCE(NULLIF(?,'')),telefono=COALESCE(NULLIF(?,'')),referente=COALESCE(NULLIF(?,'')),tipo_cliente=COALESCE(NULLIF(?,'')),settore=COALESCE(NULLIF(?,'')) WHERE id=?",
                    (row.get("VIA", ""), row.get("COMUNE", ""), row.get("CAP", ""),
                     row.get("EMAIL", ""), row.get("TELEFONO", ""), row.get("REFERENTE", ""),
                     row.get("TIPO_CLIENTE", ""), row.get("SETTORE", ""), existing["id"]),
                )
                updated += 1
            else:
                conn.execute(
                    "INSERT INTO clienti (nome_studio,via,citta,cap,email,telefono,referente,tipo_cliente,settore,note,data_inserimento) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
                    (nome, row.get("VIA", ""), row.get("COMUNE", ""), row.get("CAP", ""),
                     row.get("EMAIL", ""), row.get("TELEFONO", ""), row.get("REFERENTE", ""),
                     row.get("TIPO_CLIENTE", "Amministratore"), row.get("SETTORE", ""),
                     row.get("NOTE", ""), now),
                )
                imported += 1
        except Exception as e:
            skipped += 1

    conn.commit()
    conn.close()
    return jsonify({"ok": True, "data": {"imported": imported, "updated": updated, "skipped": skipped}})


# ── Prompt 5: Dashboard KPI ──────────────────────────────────────────

@app.route("/api/dashboard/kpi", methods=["GET"])
def api_dashboard_kpi():
    conn = get_db()
    cats = ["installazione", "servizi", "cc_modus", "cu_unitron", "fornitura", "interventi"]
    result = {}

    for cat in cats:
        if cat == "installazione":
            for sub in ["CK", "CL"]:
                key = sub
                inviate = conn.execute("SELECT COUNT(*) FROM offerte WHERE sottotipo=? AND stato_versione='attiva'", (sub,)).fetchone()[0]
                prese = conn.execute("SELECT COUNT(*) FROM offerte WHERE sottotipo=? AND stato='preso_lavoro' AND stato_versione='attiva'", (sub,)).fetchone()[0]
                val_normale = conn.execute("SELECT COALESCE(SUM(valore_commessa),0) FROM offerte WHERE sottotipo=? AND stato='preso_lavoro' AND stato_versione='attiva' AND (is_gara_appalto=0 OR is_gara_appalto IS NULL)", (sub,)).fetchone()[0]
                val_gare = conn.execute("SELECT COALESCE(SUM(vg),0) FROM (SELECT MAX(valore_gara) as vg FROM offerte WHERE sottotipo=? AND stato='preso_lavoro' AND stato_versione='attiva' AND is_gara_appalto=1 AND gara_id IS NOT NULL GROUP BY gara_id)", (sub,)).fetchone()[0]
                result[key] = {"inviate": inviate, "prese": prese, "valore_preso": val_normale + val_gare}
        elif cat == "servizi":
            inviate = conn.execute("SELECT COUNT(*) FROM offerte WHERE macro_categoria='servizi' AND stato_versione='attiva'").fetchone()[0]
            prese = conn.execute("SELECT COUNT(*) FROM offerte WHERE macro_categoria='servizi' AND stato='preso_lavoro' AND stato_versione='attiva'").fetchone()[0]
            val = conn.execute("SELECT COALESCE(SUM(importo_servizio_annuo),0) FROM offerte WHERE macro_categoria='servizi' AND stato='preso_lavoro' AND stato_versione='attiva'").fetchone()[0]
            sottotipi_attivi = [r[0] for r in conn.execute("SELECT DISTINCT sottotipo FROM offerte WHERE macro_categoria='servizi' AND sottotipo IS NOT NULL AND stato_versione='attiva'").fetchall()]
            result["servizi"] = {"inviate": inviate, "prese": prese, "valore_annuo": val, "sottotipi": sottotipi_attivi}
        else:
            inviate = conn.execute("SELECT COUNT(*) FROM offerte WHERE macro_categoria=? AND stato_versione='attiva'", (cat,)).fetchone()[0]
            prese = conn.execute("SELECT COUNT(*) FROM offerte WHERE macro_categoria=? AND stato='preso_lavoro' AND stato_versione='attiva'", (cat,)).fetchone()[0]
            val_normale = conn.execute("SELECT COALESCE(SUM(valore_commessa),0) FROM offerte WHERE macro_categoria=? AND stato='preso_lavoro' AND stato_versione='attiva' AND (is_gara_appalto=0 OR is_gara_appalto IS NULL)", (cat,)).fetchone()[0]
            val_gare = conn.execute("SELECT COALESCE(SUM(vg),0) FROM (SELECT MAX(valore_gara) as vg FROM offerte WHERE macro_categoria=? AND stato='preso_lavoro' AND stato_versione='attiva' AND is_gara_appalto=1 AND gara_id IS NOT NULL GROUP BY gara_id)", (cat,)).fetchone()[0]
            gare_inviate = conn.execute("SELECT COUNT(DISTINCT gara_id) FROM offerte WHERE macro_categoria=? AND is_gara_appalto=1 AND stato_versione='attiva'", (cat,)).fetchone()[0]
            result[cat] = {"inviate": inviate, "prese": prese, "valore_preso": val_normale + val_gare, "gare_inviate": gare_inviate}

    conn.close()
    return jsonify({"ok": True, "data": result})


# ── Startup ──────────────────────────────────────────────────────────

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
init_db()

if __name__ == "__main__":
    app.run(debug=True, port=5000, threaded=True)
